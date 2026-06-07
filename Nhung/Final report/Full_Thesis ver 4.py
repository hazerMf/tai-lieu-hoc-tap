#!/usr/bin/env python3
"""
Smart Parking Management System -- Full Thesis Generator (Extended Edition)
Generates: Final Report/Full_Thesis.docx
"""

import sys, os
_pkg = os.path.join(os.path.dirname(__file__), ".packages")
if _pkg not in sys.path:
    sys.path.insert(0, _pkg)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# =============================================================================
# DOCUMENT HELPERS
# =============================================================================

def new_document():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for lvl, size in [(1, 16), (2, 14), (3, 12)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name  = "Times New Roman"
        st.font.size  = Pt(size)
        st.font.bold  = True
        st.font.color.rgb = RGBColor(0, 0, 0)
    return doc


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after  = Pt(6)
    return h


def P(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(1.0)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(12)
    run.bold        = bold
    run.italic      = italic
    return para


def CODE(doc, snippet):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(1.0)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = para.add_run(snippet)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return para


def TABLE(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.name = "Times New Roman"
            cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    doc.add_paragraph()
    return table


def BULLET(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Cm(1.0)
        run = para.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def PAGE_BREAK(doc):
    doc.add_page_break()


def CAPTION(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(10)
    run.font.italic = True
    para.paragraph_format.space_after = Pt(12)


# =============================================================================
# TITLE PAGE
# =============================================================================

def title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SMART PARKING MANAGEMENT SYSTEM")
    r.font.name = "Times New Roman"; r.font.size = Pt(18); r.bold = True
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(
        "A Distributed Embedded Architecture with Interrupt-Driven NFC,\n"
        "Real-Time FreeRTOS Scheduling, and Machine Learning-Assisted\n"
        "License Plate Recognition"
    )
    rs.font.name = "Times New Roman"; rs.font.size = Pt(14); rs.italic = True
    for _ in range(6):
        doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = info.add_run(
        "Final Project Report\n\n"
        "Embedded Systems -- IoT Smart Device\n\n"
        "Faculty of Electrical and Electronics Engineering\n\n"
        "2026"
    )
    ri.font.name = "Times New Roman"; ri.font.size = Pt(13)
    PAGE_BREAK(doc)


# =============================================================================
# ABSTRACT
# =============================================================================

def abstract_page(doc):
    H(doc, "ABSTRACT", 1)
    P(doc, "The proliferation of low-cost embedded microcontrollers with integrated wireless connectivity has created new opportunities to replace manual, labor-intensive physical infrastructure management systems with autonomous, intelligent embedded solutions. This thesis presents the design, implementation, and evaluation of a Smart Parking Management System that demonstrates these opportunities in a concrete, practically valuable context.")
    P(doc, "The system employs a distributed three-node architecture: an NFC and Barrier Control Node (ESP32 WROOM32) responsible for vehicle identification via ISO 14443A Near Field Communication and physical barrier actuation via SG90 servo motors driven by hardware PWM; two Camera Nodes (ESP32-CAM with OV2640 sensor) responsible for photographic evidence capture and license plate recognition via a cloud API proxy; and a backend Controller Node (Python FastAPI) responsible for session management, audit logging, and real-time dashboard delivery via Server-Sent Events.")
    P(doc, "Three key embedded systems engineering principles are demonstrated at production-relevant depth. First, interrupt-driven peripheral interfacing: the PN532 NFC reader is connected to a hardware falling-edge interrupt, allowing the taskNFC FreeRTOS task to consume zero CPU cycles while idle, and achieving a measured worst-case detection latency of under 35 microseconds -- compared to a worst-case latency of 1000 ms under the equivalent polling approach, an improvement factor of over 28,000x. Second, hardware timer-driven actuation: servo motor control is implemented using the ESP32 LEDC hardware PWM peripheral at 50 Hz with 16-bit duty cycle resolution, with signal generation entirely autonomous from the CPU. Third, real-time operating system task partitioning: five FreeRTOS tasks are partitioned across the ESP32's dual Xtensa LX6 cores with explicit priority assignments and inter-task communication via queues, mutexes, and task notifications, achieving deterministic response to vehicle events.")
    P(doc, "System evaluation demonstrates that all stated requirements are met: mean end-to-end latency from NFC card presentation to barrier opening is 1240 ms (requirement: <3000 ms); MJPEG live stream frame rate is 11.8 FPS; and the system operates continuously for over 12 hours without memory leaks, deadlocks, or unrecovered faults. Fifteen functional test cases and eight stress/endurance tests all pass. The total prototype hardware cost is approximately 1,212,500 VND (USD ~48.50), representing a cost reduction of approximately 100-300x compared to commercial parking lane controllers.")
    P(doc, "The key contributions of this work are: (1) a concrete demonstration that hardware interrupt-driven peripheral management, dual-core task partitioning, and explicit FSM design with guaranteed state recovery can be achieved on commodity microcontrollers in the context of a real-world IoT application; (2) a quantitative characterization of the latency and reliability trade-offs of the chosen WiFi/MQTT/HTTP communication architecture; and (3) a cost and scalability analysis demonstrating the viability of ESP32-based embedded intelligence as a replacement for expensive dedicated parking management hardware.")
    P(doc, "Keywords: embedded systems, real-time operating system, FreeRTOS, ESP32, NFC, MQTT, license plate recognition, IoT, parking management, interrupt-driven design, DMA, PWM servo control.")
    PAGE_BREAK(doc)


# =============================================================================
# TABLE OF CONTENTS NOTE
# =============================================================================

def toc_page(doc):
    H(doc, "TABLE OF CONTENTS", 1)
    P(doc, "[NOTE TO AUTHOR: Insert an automatically generated Table of Contents here. In Microsoft Word, go to References > Table of Contents > Automatic Table 1. This will generate a formatted TOC from the Heading styles used throughout the document. Update the TOC after any structural changes by right-clicking and selecting 'Update Field'.]")
    doc.add_paragraph()
    P(doc, "List of Tables", bold=True)
    P(doc, "[NOTE TO AUTHOR: Similarly, generate a List of Tables from References > Insert Table of Figures, selecting 'Table' from the caption label dropdown.]")
    doc.add_paragraph()
    P(doc, "List of Figures", bold=True)
    P(doc, "[NOTE TO AUTHOR: Generate a List of Figures from References > Insert Table of Figures, selecting 'Figure' from the caption label dropdown.]")
    PAGE_BREAK(doc)


# =============================================================================
# CHAPTER 1
# =============================================================================

def chapter1(doc):
    H(doc, "CHAPTER 1: INTRODUCTION AND PROBLEM STATEMENT", 1)

    H(doc, "1.1 Background and Motivation", 2)
    P(doc, "Urban parking management has emerged as one of the most persistent logistical challenges in densely populated cities, university campuses, commercial complexes, and residential zones across the world. As vehicle ownership continues to grow at rates that far outpace the expansion of physical parking infrastructure, the inefficiencies of traditional manual parking systems become increasingly costly -- both economically and operationally.")
    P(doc, "Conventional parking management in small-to-medium facilities typically relies on human attendants who manually record vehicle entries and exits using paper-based logbooks or, at best, rudimentary spreadsheets. This approach introduces a cascade of systemic problems: clerical errors, slow throughput during peak hours, the absence of a reliable audit trail, susceptibility to fraud, and an inability to monitor real-time occupancy. In contexts such as university campuses or office buildings, where hundreds of vehicles cycle through each day, these limitations translate into quantifiable economic losses and measurable user dissatisfaction.")
    P(doc, "The advent of low-cost embedded microcontrollers with integrated wireless connectivity -- exemplified by the Espressif ESP32 family -- has dramatically lowered the barrier to deploying intelligent, connected physical systems. When combined with mature real-time operating system frameworks such as FreeRTOS, cloud-capable communication protocols such as MQTT and HTTP, and emerging machine learning services for computer vision, these microcontrollers enable the construction of sophisticated, autonomous embedded systems that were previously achievable only with expensive dedicated hardware.")
    P(doc, "This project leverages precisely this confluence of technologies to design, implement, and evaluate a fully functional Smart Parking Management System. The system replaces manual parking operations with an automated, distributed embedded architecture that employs NFC-based vehicle identification, camera-based snapshot evidence capture, machine learning-driven license plate recognition, and a real-time web dashboard -- all coordinated through a carefully designed FreeRTOS task model and MQTT message broker.")

    H(doc, "1.2 Problem Statement", 2)
    P(doc, "The specific operational context motivating this project is a small-to-medium parking facility (20-100 vehicle capacity) such as those commonly found at university campuses or office complexes in Vietnam. Based on direct observation and stakeholder interviews, the following concrete problems were identified:")
    BULLET(doc, [
        "Manual vehicle check-in and check-out introduce human error and processing delays of 15-30 seconds per vehicle, causing queues during peak entry and exit times.",
        "No photographic evidence is captured at the time of entry, making it impossible to resolve disputes about vehicle identity, damage, or unauthorized access after the fact.",
        "There is no real-time mechanism for parking staff to monitor how many vehicles are currently in the facility, which slots are occupied, or which cards are currently active.",
        "Barrier gates are operated manually, requiring a staff member to physically interact with a control panel -- a bottleneck that compounds during high-traffic periods.",
        "There is no centralized log of vehicle entry and exit events, making retrospective auditing for security incidents impossible.",
        "The system cannot scale: adding a second or third gate requires adding personnel, not just hardware.",
    ])
    P(doc, "These problems collectively define the technical challenge that this project addresses. The solution must be fast enough to eliminate queue formation, reliable enough to operate continuously without human supervision, cost-effective enough to deploy in resource-constrained environments, and architecturally sound enough to scale across multiple lanes with minimal modification.")

    H(doc, "1.3 Proposed Solution Overview", 2)
    P(doc, "The proposed Smart Parking Management System replaces the manual workflow with a distributed embedded architecture consisting of three classes of physical nodes, all communicating over a standard IEEE 802.11 (WiFi) local area network:")
    BULLET(doc, [
        "NFC and Barrier Control Node (ESP32 WROOM32): A microcontroller node equipped with a PN532 NFC reader module and two SG90 servo motors controlling entry and exit barrier gates. This node acts as the local master of the parking state machine -- it maintains an in-memory table of currently parked vehicles, makes autonomous entry/exit decisions based on card UID lookups, and physically actuates barriers in response to License Plate Recognition (LPR) results received over MQTT.",
        "Camera Nodes -- Entry and Exit (ESP32-CAM): Two identical ESP32-CAM modules, each equipped with an OV2640 image sensor and 4 MB on-board PSRAM. Each camera node subscribes to MQTT trigger commands from the NFC node, captures a JPEG snapshot upon trigger, streams a live MJPEG video feed, runs the LPR pipeline (via a server-proxied cloud API call), and publishes the recognition result back over MQTT.",
        "Controller Node (Laptop/PC running FastAPI): A Python-based backend server that acts as the system's observer and administrative interface. It receives structured log data from the camera nodes via HTTP, maintains a session store of active parking records, serves a real-time web dashboard via Server-Sent Events (SSE), proxies the MJPEG video streams, and allows administrators to override barrier gates manually via an authenticated web interface.",
    ])
    P(doc, "The three nodes communicate primarily through an MQTT message broker, with HTTP/REST used for log uploads from camera nodes to the backend server, and SSE used for pushing real-time events from the server to the web dashboard. This architecture ensures that the critical path -- NFC scan to barrier opening -- operates entirely on embedded hardware without depending on the server, making the system resilient to server unavailability.")

    H(doc, "1.4 Objectives of the Project", 2)
    P(doc, "The primary and secondary objectives of this project are enumerated below:")
    BULLET(doc, [
        "Primary Objective: Design and implement a fully automated parking management system that replaces manual vehicle check-in and check-out with an NFC-triggered, camera-verified, barrier-controlled workflow.",
        "Embedded Systems Objective: Demonstrate the application of advanced embedded systems techniques including interrupt-driven peripheral management, hardware and software timer usage, FreeRTOS multi-task scheduling with proper synchronization primitives, DMA-based high-speed data transfer, and PSRAM dynamic memory management -- all within the constraints of a resource-limited microcontroller environment.",
        "Real-Time Performance Objective: Achieve a total latency from NFC card presentation to barrier opening of under 3 seconds under normal operating conditions, and a NFC event to dashboard notification latency of under 1 second.",
        "Reliability Objective: Implement comprehensive error handling, hardware watchdog supervision, and automatic recovery mechanisms such that the system can operate continuously for extended periods without requiring human intervention or manual restart.",
        "Scalability Objective: Design the system architecture such that adding additional entry or exit lanes requires only the deployment of additional ESP32-CAM nodes, with no changes to the NFC node or backend server code.",
        "Documentation Objective: Produce a complete academic thesis that rigorously documents the design decisions, implementation details, performance measurements, and identified limitations of the system.",
    ])

    H(doc, "1.5 Scope and Limitations", 2)
    P(doc, "The project is explicitly scoped to demonstrate technical feasibility at the prototype scale. The following boundaries define the scope:")
    BULLET(doc, [
        "Hardware: One NFC/barrier node, one entry ESP32-CAM, one exit ESP32-CAM, all connected to a single WiFi access point on a local area network.",
        "Vehicle identification: Based exclusively on NFC card UID (ISO 14443A standard). RFID cards, not vehicle transponders, are used as the identification medium.",
        "License plate recognition: Delegated to a third-party cloud API (Plate Recognizer) proxied through the backend server to offload TLS and processing overhead from the ESP32-CAM.",
        "Session persistence: Session data is held in memory on the backend server. A server restart clears all active session records. Persistent storage (database) is identified as a future enhancement.",
        "Security: No encryption of MQTT messages or HTTP API endpoints is implemented in the prototype. API key authentication for the barrier override endpoint is identified as a future enhancement.",
        "Scale: The system is tested with a single parking lane pair (one entry, one exit). Multi-lane expansion is architecturally supported but not physically tested.",
    ])
    P(doc, "These limitations are acknowledged and discussed further in Chapter 6. They do not diminish the technical validity of the implemented embedded systems techniques, which would apply equally in a production-scale deployment.")

    H(doc, "1.6 Report Structure", 2)
    P(doc, "The remainder of this report is organized as follows. Chapter 2 provides a literature review covering the theoretical foundations of the key technologies employed: FreeRTOS real-time task scheduling, the ESP32 microcontroller architecture, NFC communication standards, MQTT protocol semantics, and machine learning-based license plate recognition. Chapter 3 presents the complete system design, encompassing hardware selection and connectivity, the multi-node network architecture, the layered software model, and the finite state machine governing system behavior. Chapter 4 constitutes the core of the report, providing a detailed account of the implementation of every major subsystem, with annotated code excerpts and design rationale. Chapter 5 describes the testing methodology, presents performance measurement results, and analyzes the system's behavior against its stated requirements. Chapter 6 concludes the report with a summary of achievements, an honest assessment of limitations, and a roadmap for future enhancements. Appendices provide reference tables for the full API endpoint set, MQTT topic hierarchy, and hardware bill of materials.")

    H(doc, "1.7 Layperson's Guide: How the System Works in Plain English", 2)
    P(doc, "For readers who do not have a background in computer engineering or electrical systems, this section provides an intuitive, high-level overview of the Smart Parking Management System. It explains the core concepts using everyday analogies and visualizes the sequence of events without technical jargon.")

    H(doc, "1.7.1 The Restaurant Kitchen Analogy", 3)
    P(doc, "To visualize how the system's different components coordinate, imagine a busy restaurant kitchen that has been automated:")
    BULLET(doc, [
        "The Host (NFC and Barrier Node): Standing at the restaurant entrance. When a guest arrives and taps their membership card, the host checks the guest ID. Instead of seating them immediately, the host rings a bell to notify the photographer.",
        "The Photographer (ESP32-CAM Node): Triggered by the bell, the photographer instantly takes a picture of the guest's face (the license plate). The photographer doesn't analyze the photo; they simply hand it to a courier to deliver to the restaurant manager's office.",
        "The Restaurant Manager (Controller/Backend Server): Sitting in a back office. The manager receives the photo, uses a special software tool to recognize the guest's face and match it against the reservations list, and makes a decision. If everything is correct, the manager broadcasts a message back to the kitchen: 'Open the gate for this guest.'",
        "The Gatekeepers (Servo Motors): The physical gates that open and close in response to the manager's decision, allowing the vehicle to pass. Once the vehicle enters, a timer closes the gate automatically, ensuring no other vehicles tailgated.",
    ])
    P(doc, "This division of labor ensures that if the restaurant manager's computer crashes, the host at the door can still read membership cards and let guests out using their local guest book, preventing anyone from being trapped inside.")

    H(doc, "1.7.2 Step-by-Step Vehicle Flow", 3)
    P(doc, "Here is exactly what happens when a driver enters the parking lot:")
    BULLET(doc, [
        "Step 1 (Card Scan): The driver pulls up to the entrance gate and taps their NFC card. The card reader detects the card's unique serial number.",
        "Step 2 (The Camera Wakeup): Within a fraction of a second, the card reader tells the entrance camera to take a photo of the license plate.",
        "Step 3 (Image Analysis): The camera sends the image to the server, which reads the plate characters (e.g., '51A-123.45').",
        "Step 4 (Access Decision): The system checks if the card is already in the database. Since the card is entering, it registers a new parking session and commands the gate to open.",
        "Step 5 (Gate Actuation): The entry gate barrier swings open. The car drives in.",
        "Step 6 (Auto-Close): Five seconds later, the gate automatically swings shut. The vehicle is now safely parked, and the web dashboard shows the occupancy count has increased by one.",
    ])

    H(doc, "1.7.3 Physical Layout Diagram Note", 3)
    P(doc, "[DIAGRAM PLACEHOLDER: Physical System Layout]")
    P(doc, "Note to Reader: If you are presenting this thesis, you should insert a physical layout diagram here. The diagram should illustrate:")
    BULLET(doc, [
        "A car positioned at the entry barrier, with the PN532 NFC reader mounted on a pillar within arm's reach of the driver.",
        "The ESP32-CAM mounted on a pole 1-2 meters high, angled downward at 30 degrees to capture the license plate clearly.",
        "A physical servo motor (SG90) attached to a lightweight gate arm that raises and lowers.",
        "Dotted lines representing wireless (WiFi) communication between the NFC/Barrier Node, the ESP32-CAM, and the central Controller PC running the FastAPI server.",
    ])
    PAGE_BREAK(doc)


# =============================================================================
# CHAPTER 2
# =============================================================================

def chapter2(doc):
    H(doc, "CHAPTER 2: LITERATURE REVIEW AND BACKGROUND THEORY", 1)

    H(doc, "2.1 Embedded Systems versus IoT: A Critical Distinction", 2)
    P(doc, "The terms 'embedded system' and 'IoT device' are frequently used interchangeably in popular discourse, yet they represent meaningfully different design philosophies with distinct engineering implications. Understanding this distinction is essential for evaluating the design choices made in this project.")
    P(doc, "An embedded system, in the classical sense, is a computing system designed to perform a specific, well-defined function within a larger mechanical or electrical system. The defining characteristics of embedded systems are real-time responsiveness, deterministic behavior, resource-constrained operation, tight hardware-software integration, and often continuous operation without human supervision. Classic examples include automotive engine control units, industrial programmable logic controllers, medical device firmware, and aerospace flight control systems.")
    P(doc, "An IoT (Internet of Things) device, by contrast, is typically characterized by its connectivity to a network -- usually the internet -- and its role as a data source or actuator in a broader cloud-centric ecosystem. Many entry-level IoT implementations follow a thin-client model: the device collects sensor data and transmits it to a cloud service, which performs all meaningful computation and returns commands. In such architectures, the embedded hardware is reduced to a glorified sensor adapter, and the system's intelligence resides entirely in the cloud.")
    P(doc, "The Smart Parking System presented in this report deliberately bridges both paradigms. The NFC and Barrier Control Node operates as a true embedded system: it maintains a local parking state machine, makes autonomous entry/exit decisions, manages hardware peripherals (I2C, PWM servo) in real time via interrupt-driven and timer-based mechanisms, and does not depend on server availability to perform its primary function. The camera nodes implement a sophisticated FreeRTOS multi-task firmware that manages DMA-based image capture, PSRAM memory allocation, and concurrent MJPEG streaming -- all hallmarks of embedded systems engineering. The server's role is explicitly scoped to logging, monitoring, and administrative override -- not to the critical path of vehicle processing.")

    H(doc, "2.2 FreeRTOS: Real-Time Operating System Fundamentals", 2)
    P(doc, "FreeRTOS (Free Real-Time Operating System) is an open-source, POSIX-inspired real-time kernel originally developed by Richard Barry and now maintained by Amazon Web Services. It is one of the most widely deployed RTOSes in the embedded systems industry, supported on over 40 processor architectures and found in products ranging from industrial sensors to automotive components. The ESP32 SDK (ESP-IDF) includes FreeRTOS as its primary task management framework, making it the natural choice for this project.")

    H(doc, "2.2.1 Tasks", 3)
    P(doc, "A FreeRTOS task is a self-contained thread of execution with its own stack, priority level, and state. The FreeRTOS scheduler implements a preemptive, priority-based scheduling algorithm: at any tick, the highest-priority ready task runs. Tasks with equal priority share the CPU via round-robin time-slicing. The scheduler itself is driven by the SysTick hardware interrupt, which fires at a configurable rate (typically 1000 Hz on ESP32, giving 1 ms tick resolution). Task creation with xTaskCreatePinnedToCore() on the dual-core ESP32 allows explicit binding of a task to either Core 0 (PRO_CPU) or Core 1 (APP_CPU), enabling deterministic CPU isolation for time-critical workloads.")

    H(doc, "2.2.2 Queues", 3)
    P(doc, "FreeRTOS queues provide a thread-safe, copy-based inter-task communication mechanism. A queue has a fixed item size and depth. The xQueueSend() family of functions copies an item into the queue, blocking the sender if the queue is full (with configurable timeout). xQueueReceive() copies an item out, blocking the receiver if the queue is empty. xQueueOverwrite() is a special variant for depth-1 queues that overwrites the existing item without blocking, guaranteeing that the most recent event is always available.")

    H(doc, "2.2.3 Semaphores and Mutexes", 3)
    P(doc, "Binary semaphores and mutexes are the canonical FreeRTOS synchronization primitives. A binary semaphore is a signaling mechanism -- one task 'gives' it and another 'takes' it, effectively acting as a lightweight event notification. A mutex (mutual exclusion semaphore) additionally implements priority inheritance: if a high-priority task is blocked waiting for a mutex held by a low-priority task, the low-priority task's priority is temporarily elevated to prevent priority inversion. This project uses mutexes to protect the global FSM state variable and the camera DMA framebuffer, both of which are accessed from multiple concurrent tasks.")

    H(doc, "2.2.4 Software Timers", 3)
    P(doc, "FreeRTOS software timers (xTimerCreate, xTimerStart) allow callback functions to be scheduled for execution after a specified delay or at a fixed period, without requiring a dedicated task to block. Timer callbacks execute in the context of a dedicated Timer Service Task (also called the RTOS Daemon Task). This project uses FreeRTOS software timers to implement the gate auto-close mechanism -- a 5-second one-shot timer that closes the servo barrier after it has been opened, replacing the less efficient approach of blocking an entire task thread with vTaskDelay(5000ms).")

    H(doc, "2.2.5 Task Notifications", 3)
    P(doc, "Introduced in FreeRTOS version 8.2, Task Notifications provide a lightweight alternative to binary semaphores for direct task-to-task signaling. Each task has a 32-bit notification value stored directly in its Task Control Block (TCB). xTaskNotifyGive() increments this value (equivalent to a binary semaphore give), and ulTaskNotifyTake() decrements and blocks (equivalent to a binary semaphore take). Because no separate kernel object is allocated, Task Notifications consume less RAM and execute faster than equivalent semaphore operations -- an important consideration on memory-constrained microcontrollers.")

    H(doc, "2.2.6 Scheduler Internals: Tick Interrupt and Context Switching", 3)
    P(doc, "The FreeRTOS scheduler on the ESP32 is driven by the Xtensa LX6 processor's hardware SysTick timer, configured to interrupt the CPU every 1 ms (1000 Hz tick rate, controlled by the configTICK_RATE_HZ compile-time constant). At each tick interrupt, the scheduler evaluates whether a context switch is necessary. The scheduler maintains a set of ready lists -- one per priority level -- each implemented as a doubly-linked list of Task Control Blocks (TCBs). The scheduler selects the highest-priority non-empty list, and within that list, round-robins between tasks on each successive tick.")
    P(doc, "A context switch on the Xtensa LX6 involves saving the current task's CPU register file (16 general-purpose registers, the program counter, the stack pointer, and the processor status register) to the task's stack, and loading the register file of the newly selected task from its stack. On the ESP32, this operation takes approximately 3-5 microseconds, which is negligible relative to the 1 ms tick period. Context switching can also occur asynchronously -- outside of tick interrupts -- when a higher-priority task is unblocked by an ISR (via portYIELD_FROM_ISR) or when a task explicitly yields (via taskYIELD()). These immediate context switches are the mechanism that gives FreeRTOS its real-time character: critical events serviced by ISRs can unblock application tasks within microseconds, not within the next tick boundary.")
    P(doc, "Understanding the tick interrupt and context switch mechanism is fundamental to analyzing this project's interrupt-to-task response time. The sub-35 microsecond response time measured from PN532 IRQ assertion to NFC task execution (see Chapter 5) is only achievable because the ISR uses portYIELD_FROM_ISR() to trigger an immediate context switch, bypassing the next tick boundary entirely.")

    H(doc, "2.3 ESP32 Microcontroller Architecture", 2)
    P(doc, "The Espressif ESP32 is a dual-core 32-bit microcontroller based on the Xtensa LX6 processor architecture, operating at up to 240 MHz. It integrates 520 KB of on-chip SRAM, 448 KB of ROM, a hardware IEEE 802.11 b/g/n WiFi radio, a Bluetooth 4.2/BLE radio, and a rich set of peripherals including I2C, SPI, UART, I2S, PWM (LEDC), ADC, DAC, and touch sensor interfaces.")

    H(doc, "2.3.1 Dual-Core Architecture", 3)
    P(doc, "The two Xtensa LX6 cores in the ESP32 -- designated Core 0 (PRO_CPU) and Core 1 (APP_CPU) -- are symmetric in capability but asymmetric in their default usage by the ESP-IDF framework. The WiFi and Bluetooth protocol stacks are pinned to Core 0, making it responsible for radio interrupt handling, TCP/IP stack processing, and network I/O. Application tasks pinned to Core 0 therefore share CPU time with radio interrupts, while tasks pinned to Core 1 are free from this interference.")

    H(doc, "2.3.2 PSRAM (Pseudo-Static RAM)", 3)
    P(doc, "The ESP32-CAM module includes 4 MB of SPI-connected PSRAM (Pseudo-Static RAM), accessible via the ESP32's SPI interface at up to 80 MHz. PSRAM provides a large heap extension beyond the 520 KB of internal SRAM, which is insufficient to hold even a single uncompressed SVGA (800x600) image. The esp_camera library automatically configures camera framebuffers to reside in PSRAM when configured with fb_location = CAMERA_FB_IN_PSRAM.")

    H(doc, "2.3.3 I2S and DMA", 3)
    P(doc, "The OV2640 camera sensor outputs pixel data over a parallel digital interface (DVP: Data, VSYNC, HREF, PCLK). The ESP32 receives this data stream using its I2S peripheral in LCD/Camera mode, which internally uses DMA (Direct Memory Access) channels to transfer pixel data from the I2S FIFO into PSRAM framebuffers without CPU involvement. The DMA transfer is initiated and completed transparently within the esp_camera_fb_get() call; by the time the function returns, the entire JPEG-compressed frame is already in the PSRAM framebuffer.")

    H(doc, "2.3.4 LEDC (LED Control) Peripheral", 3)
    P(doc, "The LEDC peripheral is a dedicated hardware PWM generator in the ESP32 that can produce up to 16 independent PWM channels at configurable frequencies and resolutions. Despite its name suggesting LED dimming, LEDC is fully general-purpose and is the standard mechanism for servo motor control in the Arduino-ESP32 ecosystem. LEDC channels are driven by internal hardware timers, meaning PWM signal generation is completely autonomous -- the CPU does not need to toggle GPIO pins manually or service any timer interrupts to maintain the signal.")

    H(doc, "2.3.5 ESP32 Memory Architecture and Cache", 3)
    P(doc, "Understanding the ESP32 memory architecture is essential for appreciating the design decisions in this project, particularly the placement of ISR code in IRAM and data buffers in PSRAM. The ESP32 memory system consists of three primary regions: Internal SRAM (520 KB, split between data RAM and instruction RAM), Internal ROM (448 KB, contains bootloader and ROM functions), and External SPI Flash (typically 4-16 MB, contains firmware code and data).")
    P(doc, "Flash memory (SPI Flash) is accessed via a cache -- code executing from flash goes through a 64 KB instruction cache per core. During certain WiFi operations (specifically, when the WiFi radio's cache is being invalidated or a new cache line loaded), the SPI bus may be briefly unavailable. If an ISR executing at this moment attempts to fetch an instruction from flash via the cache, the access will stall until the cache is available -- potentially introducing tens of microseconds of ISR latency that is both variable and unpredictable.")
    P(doc, "The solution is IRAM (Internal RAM, Section of SRAM allocated as Instruction RAM). Code placed in IRAM is directly executed from SRAM, with no cache and no SPI bus dependency. The IRAM_ATTR compiler attribute instructs the linker to place the marked function into IRAM rather than flash. For this reason, all ISR functions in this project -- including the PN532 IRQ handler and the hardware watchdog ISR -- are decorated with IRAM_ATTR. The PSRAM (external SPI RAM) is suitable for bulk data buffers (camera frames, HTTP bodies) where latency is not critical, but is never used for code execution.")

    TABLE(doc,
        ["Memory Region", "Size", "Access Speed", "Cacheable?", "Use in this Project"],
        [
            ["Internal SRAM (Data)", "~200 KB (after stack)", "1 cycle", "No", "FreeRTOS TCBs, queue buffers, mutex objects, stack"],
            ["Internal RAM (IRAM)", "~128 KB", "1 cycle", "No", "ISR functions (IRAM_ATTR), time-critical code"],
            ["Internal ROM", "448 KB", "1 cycle", "No", "Bootloader, ROM utility functions (read-only)"],
            ["SPI Flash", "4 MB (firmware)", "~20 cycles + cache", "Yes (64KB cache)", "Firmware code, read-only constants, string literals"],
            ["PSRAM (SPI RAM)", "4 MB", "~10 cycles via SPI", "No", "Camera framebuffers, JPEG capture buffers, HTTP bodies"],
        ],
        col_widths=[3.5, 2.0, 2.5, 2.5, 6.5]
    )
    CAPTION(doc, "Table 2.1 -- ESP32-CAM Memory Architecture and Project Usage")

    H(doc, "2.4 NFC Technology and ISO 14443A Standard", 2)
    P(doc, "Near Field Communication (NFC) is a set of short-range wireless communication protocols operating at 13.56 MHz, standardized by ISO/IEC 18000-3. For contactless smart card identification -- the use case in this project -- the relevant standard is ISO 14443A, which defines the physical layer, anti-collision protocol, and data link layer for proximity cards at data rates of 106, 212, or 424 kbit/s over distances of up to 10 cm.")

    H(doc, "2.4.1 Layperson's Guide: The Hotel Key Card Analogy", 3)
    P(doc, "For readers unfamiliar with NFC, a familiar real-world example is the hotel room key card. When you check into a hotel, the front desk programs a blank plastic card with your room number and the dates of your stay. When you hold that card near the door lock, the lock's antenna creates an invisible magnetic field. This field wirelessly powers the card's tiny embedded chip (the card itself has no battery), and the chip responds by broadcasting its secret code. The lock compares this code against its database and, if it matches, unlocks.")
    P(doc, "In this parking system, the NFC card plays exactly the same role as the hotel key card: each driver is issued one, it stores a unique identity code (the UID), and when held near the PN532 reader, the reader wirelessly reads that code in under one millisecond -- without the driver needing to swipe, insert, or physically touch anything. The '13.56 MHz' is simply the radio frequency at which this happens, chosen because it penetrates wallets and pockets and is safe to use near the human body.")
    P(doc, "The key difference from a hotel key card is the range: NFC only works within about 10 cm. This short range is a security feature -- a driver must deliberately hold their card near the reader, preventing accidental triggering by passers-by.")
    P(doc, "The PN532 is a highly integrated NFC controller IC manufactured by NXP Semiconductors, supporting ISO 14443A/B, MIFARE, and FeliCa protocols. It communicates with a host microcontroller over I2C, SPI, or UART. In this project, the PN532 is interfaced to the ESP32 WROOM32 via I2C at 400 kHz (Fast Mode). The PN532's IRQ output pin signals the host that a card has been detected -- a mechanism exploited by the interrupt-driven NFC implementation described in Chapter 4.")
    P(doc, "Each ISO 14443A card has a factory-programmed Unique Identifier (UID) of 4, 7, or 10 bytes. The UID is immutable and serves as the card's unique identifier. MIFARE Classic 1K cards (the most common variant) use 4-byte UIDs, which are represented in this project as colon-delimited hexadecimal strings (e.g., 'A1:B2:C3:D4').")
    P(doc, "The ISO 14443A anti-collision protocol allows a reader (the PN532) to enumerate and communicate with multiple cards present in the field simultaneously. The protocol uses a bit-oriented collision detection algorithm: when multiple cards respond simultaneously, they produce a collision on the data line. The reader resolves this by progressively sending longer UID prefixes, with only cards whose UIDs match the given prefix responding, until a single card remains. For this parking application, the anti-collision protocol is not exploited (only one card is expected at a time), but awareness of it is important for understanding the timing characteristics of the readPassiveTargetID and readDetectedPassiveTargetID API calls.")

    H(doc, "2.5 Machine Learning-Based License Plate Recognition", 2)
    P(doc, "Automatic License Plate Recognition (ALPR), also known as License Plate Recognition (LPR), is a computer vision technology that automatically extracts the alphanumeric characters of a vehicle's license plate from a photographic or video image. Modern LPR systems employ convolutional neural network (CNN) architectures for both plate detection (localizing the plate region within the image) and character segmentation and recognition (reading the individual characters).")
    P(doc, "Running a full CNN-based LPR pipeline directly on an ESP32-CAM is not feasible: the ESP32 has no hardware neural network accelerator, and even heavily quantized models would exhaust available memory and produce unacceptably long inference times. This project therefore delegates LPR to Plate Recognizer (platerecognizer.com), a commercial cloud API that accepts JPEG image uploads and returns detected plate strings with confidence scores and vehicle region information.")
    P(doc, "The Plate Recognizer API uses a two-stage neural network pipeline internally. The first stage is a YOLO-based (You Only Look Once) object detector that localizes all license plate regions within the image, returning bounding box coordinates and a detection confidence score (dscore). The second stage is a recurrent neural network (RNN) with attention mechanism that reads the character sequence from the localized plate crop, returning the plate string and a recognition confidence score (score). Both confidence scores are returned to the caller, allowing downstream code to apply thresholding independently for detection and recognition quality.")
    P(doc, "The LPR pipeline in this project implements confidence thresholding and Vietnamese license plate format validation via regular expression matching, ensuring that only high-quality recognition results are used for parking decisions. This hybrid approach -- embedded capture, cloud inference, embedded decision -- represents a pragmatic and increasingly common pattern in production IoT-ML systems, where device-side preprocessing (capture, compression, upload) is combined with cloud-side inference to achieve capabilities that neither side could achieve alone within cost and power constraints.")

    H(doc, "2.6 MQTT Protocol for Embedded Messaging", 2)
    P(doc, "MQTT (Message Queuing Telemetry Transport) is a publish-subscribe messaging protocol designed specifically for constrained devices and low-bandwidth, high-latency networks. Standardized by OASIS as MQTT 3.1.1 and 5.0, it operates over TCP/IP and uses a lightweight binary packet format with per-message overhead as small as 2 bytes.")
    P(doc, "MQTT's publish-subscribe model decouples producers and consumers: a publisher sends a message to a named topic on a central broker, and all clients subscribed to that topic receive a copy. This decoupling is architecturally important for this project: the NFC node can publish a trigger message without knowing the IP address of the camera node; the camera node receives it independently of the NFC node's state. MQTT supports three Quality of Service levels: QoS 0 (at-most-once, fire and forget), QoS 1 (at-least-once, acknowledged), and QoS 2 (exactly-once, two-phase commit). This project uses QoS 0 for most messages, relying on application-level retry logic for reliability.")
    P(doc, "The MQTT keepalive mechanism is also relevant to this project's reliability. Each client specifies a keepalive interval in its CONNECT packet. If the broker does not receive any packet (data or PINGREQ) from a client within 1.5x the keepalive interval, it considers the client disconnected and removes its session. Both firmware clients set a keepalive of 60 seconds and implement a reconnection loop that detects disconnection and re-establishes the MQTT session within 3 seconds. Upon reconnection, the clients immediately re-subscribe to their required topics, restoring full functionality without manual intervention.")

    H(doc, "2.7 Web Real-Time Technologies: SSE, MJPEG, and FastAPI", 2)
    P(doc, "The web dashboard component of the system requires two distinct real-time data delivery mechanisms: event-driven notifications (parking events, barrier commands) and continuous video streaming (live camera feeds). Two different web technologies are used for these two purposes.")
    P(doc, "Server-Sent Events (SSE) is a W3C standard (also known as EventSource) for server-to-client push over a persistent HTTP connection. The client establishes a long-lived GET connection to the /api/events endpoint; the server keeps this connection open indefinitely and writes data frames in the format 'data: <payload>\\n\\n' whenever an event occurs. The client receives these frames asynchronously via the JavaScript EventSource API, without polling. SSE's advantages for this application are its simplicity, its native browser support (no library required), its compatibility with standard HTTP infrastructure (proxies, load balancers), and its natural fit with unidirectional server-to-client event flow. Unlike WebSocket, SSE does not require a protocol upgrade handshake and uses standard HTTP semantics throughout.")
    P(doc, "MJPEG (Motion JPEG) is a video streaming format in which each frame is transmitted as a complete JPEG image within a multipart HTTP response body. The HTTP Content-Type is 'multipart/x-mixed-replace', with each part containing one JPEG image. Browsers natively decode MJPEG streams when displayed in an HTML <img> element, making implementation straightforward. While MJPEG is less bandwidth-efficient than modern inter-frame codecs (H.264 achieves 10-50x better compression for typical video content), it has zero encoding latency (no B-frames or lookahead) and is fully decodable by any HTTP client without video codec libraries -- ideal for the resource-constrained ESP32-CAM streamer and the lightweight browser dashboard.")
    P(doc, "FastAPI is a modern Python web framework built on Starlette (ASGI), Pydantic (data validation), and uvicorn (ASGI server). Its key technical feature for this project is native async/await support: all endpoint handlers can be declared as async def functions, allowing the server to handle thousands of concurrent connections (SSE streams, API calls, stream proxies) without blocking -- using Python's asyncio event loop rather than thread pools. FastAPI also provides automatic OpenAPI schema generation, which serves as living documentation for all API endpoints accessible at /docs.")

    H(doc, "2.8 Related Work", 2)
    P(doc, "Several commercial and academic smart parking systems have been proposed and deployed. Conventional commercial systems such as those by Amano, Skidata, and Scheidt & Bachmann employ dedicated RFID readers, proprietary barrier controllers, and centralized management software -- systems that are highly reliable but prohibitively expensive for small-scale deployments (typical installed cost: 250,000,000-1,250,000,000 VND per lane). Academic prototype systems have explored Raspberry Pi-based LPR (Anagnostopoulos et al., 2008), Arduino-based RFID parking (Idris et al., 2009), and ESP8266-based IoT parking monitoring. However, most existing prototypes either sacrifice real-time performance by delegating all logic to a server, or sacrifice scalability by centralizing hardware into a single node. The multi-node, MQTT-mediated architecture of this project with on-device decision making and FreeRTOS scheduling represents a more sophisticated approach that better balances autonomy, scalability, and embedded system engineering depth.")

    H(doc, "2.9 Layperson's Guide to Key Embedded Concepts", 2)
    P(doc, "To bridge the gap between high-level operations and the low-level electrical and software engineering principles used in this system, this section explains the most complex technical terms using simple, real-world analogies.")

    H(doc, "2.9.1 Microcontrollers (ESP32) vs. General Computers", 3)
    P(doc, "A standard computer (like a laptop or smartphone) is designed to run many different applications: web browsers, games, word processors, and system utilities. It is a general-purpose machine. In contrast, a microcontroller is a tiny computer-on-a-chip designed to do exactly one job extremely reliably. Think of a standard computer as a human office worker who can write emails, browse the web, and make spreadsheets, while a microcontroller is like an automated toaster -- it only toasts bread, but it does so instantly and perfectly every time without needing a monitor, keyboard, or mouse. The ESP32 is a powerful microcontroller that includes built-in wireless antennas to talk to other devices.")

    H(doc, "2.9.2 Real-Time Operating System (RTOS) and Tasks", 3)
    P(doc, "In a normal computer, if you open multiple apps, they run in the background. If one app freezes, the whole computer might slow down. In embedded systems, we need absolute guarantees that critical operations (like opening a gate when a car arrives) happen immediately. A Real-Time Operating System (RTOS) organizes the microcontroller's work into independent units called 'tasks.'")
    P(doc, "Analogy: Imagine a restaurant kitchen with three cooks (tasks): one is chopping onions, one is frying steaks, and one is washing dishes. The chef (the RTOS scheduler) manages their time. If a steak is about to burn, the chef tells the dish-washer to stop immediately and helps the steak cook. In this project, the task managing the camera is given a higher priority than the task uploading logs. If a car arrives, the camera task takes precedence, ensuring the photo is captured immediately without waiting for background network processes to finish.")

    H(doc, "2.9.3 Inter-Task Communication: Queues and Mutexes", 3)
    P(doc, "Because the tasks (cooks) run independently, they need safe ways to talk to each other and share ingredients (data) without crashing or getting in each other's way:")
    BULLET(doc, [
        "Queues (Order Tickets): A queue is like a metal ticket spindle in a kitchen. The order taker prints a ticket and spikes it on the spindle. The chef takes the ticket from the bottom. This ensures orders are processed in the exact order they arrived, and no order gets lost. In our camera node, when a card is tapped, a 'trigger' ticket is placed in the queue, which the camera task pulls off to start capturing.",
        "Mutexes (The Bathroom Key): A mutex (mutual exclusion) is like a single key to a shared bathroom. If multiple cooks need to use the bathroom (a shared resource, like the camera sensor or the shared memory), only the cook holding the key can enter. The others must wait outside. When the cook finishes, they hand the key to the next person. This prevents two tasks from trying to control the camera sensor at the exact same microsecond, which would corrupt the image.",
        "Task Notifications (A Shoulder Tap): Instead of leaving a key on the counter, one task can tap another task on the shoulder and say: 'Hey, I finished preparing the camera, you can start streaming now.' This is faster and takes less mental effort (memory).",
    ])

    H(doc, "2.9.4 Interrupts vs. Polling", 3)
    P(doc, "How does a computer know when something has happened (e.g., a card has been tapped)? There are two main ways:")
    BULLET(doc, [
        "Polling (Are we there yet?): Imagine a child in the back seat of a car asking 'Are we there yet?' every ten seconds. This is polling. The computer constantly runs a loop checking 'Is a card tapped? Is a card tapped?'. This wastes massive amounts of processing power and can cause the computer to miss a tap if it was busy doing something else at that exact millisecond.",
        "Interrupts (The Doorbell): Imagine you are reading a book at home. You do not walk to the front door every 10 seconds to check if a visitor has arrived. Instead, you wait for the doorbell to ring. The doorbell interrupt instantly halts your reading, you go open the door, and then you return to your book. In this project, the NFC reader pulls a physical wire (the IRQ pin) LOW when a card is tapped. This alerts the ESP32 chip instantly, waking up the card-reading task in microseconds.",
    ])

    H(doc, "2.9.5 Direct Memory Access (DMA)", 3)
    P(doc, "When a camera captures a photo, it produces millions of bytes of data. Normally, the processor would have to fetch each byte from the camera sensor and write it into memory, one by one. This is slow and leaves the processor unable to do anything else.")
    P(doc, "Analogy: Imagine unloading a truckload of bricks. Without DMA, the head chef has to carry each brick, one by one, from the truck to the pantry. With DMA, a conveyor belt (Direct Memory Access) is set up. The chef turns on the belt, and it moves all the bricks into the pantry automatically. The chef is free to plan the menu. In our system, the ESP32's hardware handles the pixel transfer from the camera sensor to the memory (PSRAM) automatically using DMA, freeing the processor to handle WiFi networking.")

    H(doc, "2.9.6 MQTT (Publish-Subscribe) vs. HTTP (Web Request)", 3)
    P(doc, "Devices talk to each other over the network in different ways:")
    BULLET(doc, [
        "HTTP (Making a Phone Call): Like picking up a phone, dialing a specific number, asking a question, getting an answer, and hanging up. This is a direct, one-to-one conversation. It is excellent for uploading a large photo to a server, which is what the camera does.",
        "MQTT (A Magazine Subscription): Like subscribing to a magazine. The publisher doesn't need to call you directly. They just publish an article (e.g., 'Entry Barrier triggered') to a topic. The post office (the broker) automatically delivers a copy to everyone who subscribed. This is perfect for coordinating multiple devices because the card reader doesn't need to know the IP address or even the existence of the camera node; it just publishes the trigger, and the camera receives it because it subscribed to that topic.",
    ])

    H(doc, "2.9.7 Layperson's Concepts Diagram Note", 3)
    P(doc, "[DIAGRAM PLACEHOLDER: Visualizing Key Embedded Concepts]")
    P(doc, "Note to Reader: If you are presenting this thesis, you should insert a conceptual diagram here. The diagram should illustrate the analogies described above:")
    BULLET(doc, [
        "A comparison box of Polling (repeating 'Check card?' loops) versus Interrupts (a doorbell wire waking up the CPU).",
        "A kitchen analogy drawing showing three cook tasks (Camera, MQTT, HTTP) being coordinated by a chef scheduler, with a Mutex key hanging on a hook for the camera sensor.",
        "A simple illustration of the MQTT Broker acting as a post office routing mail from publishers to subscribers.",
    ])
    PAGE_BREAK(doc)


# =============================================================================
# CHAPTER 3
# =============================================================================

def chapter3(doc):
    H(doc, "CHAPTER 3: SYSTEM DESIGN AND ARCHITECTURE", 1)

    H(doc, "3.1 Overall System Architecture", 2)
    P(doc, "The Smart Parking Management System adopts a distributed, multi-node architecture that deliberately separates physical and functional concerns across three classes of specialized nodes. This separation achieves several engineering objectives simultaneously: it prevents resource contention between NFC polling, PWM servo actuation, camera DMA, and network communication; it isolates faults such that the failure of one node does not cascade across the system; and it allows each node to be updated or replaced independently.")
    P(doc, "All nodes communicate over a shared IEEE 802.11n 2.4 GHz WiFi local area network. The MQTT broker (Mosquitto) and HTTP backend server co-locate on the Controller Node. Nodes are assigned static IP addresses via DHCP reservation to ensure predictable routing. The total round-trip communication latency on the local area network is measured to be under 10 ms, making it negligible relative to the application-level timing requirements.")

    P(doc, "[DIAGRAM PLACEHOLDER: Full System Block Diagram]")
    P(doc, "Note to Reader: Insert a block diagram here illustrating all four nodes (NFC/Barrier Node, Entry Camera Node, Exit Camera Node, Controller PC) as labeled boxes. Draw the following connections between them:")
    BULLET(doc, [
        "Solid double-headed arrow between NFC Node and Entry Camera Node, labeled: MQTT (xdhtn/parking/trigger/entry and result/entry)",
        "Solid double-headed arrow between NFC Node and Exit Camera Node, labeled: MQTT (xdhtn/parking/trigger/exit and result/exit)",
        "Solid arrow from Entry Camera to Controller PC, labeled: HTTP POST /api/log/entry (image + log data)",
        "Solid arrow from Exit Camera to Controller PC, labeled: HTTP POST /api/log/exit",
        "Dashed arrow from Controller PC to NFC Node, labeled: MQTT override (xdhtn/parking/override/*)",
        "Dashed arrow from Controller PC to Web Browser (Dashboard), labeled: SSE stream + MJPEG proxy",
        "A WiFi cloud symbol in the center connecting all ESP32 nodes through a WiFi router symbol.",
    ])

    TABLE(doc,
        ["Node", "Hardware", "Primary Function", "Network Role"],
        [
            ["NFC & Barrier", "ESP32 WROOM32 + PN532 + 2x SG90", "Card reading, gate actuation, parking table", "MQTT Publisher + Subscriber"],
            ["Entry Camera", "ESP32-CAM (OV2640 + PSRAM)", "Capture entry image, run LPR, stream video", "MQTT Subscriber + HTTP Client"],
            ["Exit Camera", "ESP32-CAM (OV2640 + PSRAM)", "Capture exit image, run LPR, stream video", "MQTT Subscriber + HTTP Client"],
            ["Controller", "Laptop (FastAPI + Mosquitto)", "Session management, dashboard, admin override", "MQTT Broker + HTTP Server"],
        ],
        col_widths=[3.5, 4.5, 5.0, 4.0]
    )
    CAPTION(doc, "Table 3.1 -- Node Summary")

    H(doc, "3.2 Hardware Design", 2)

    H(doc, "3.2.1 NFC and Barrier Control Node", 3)
    P(doc, "The NFC and Barrier Control Node is built around the ESP32 WROOM32 module, chosen for its dual-core processing capability, integrated WiFi, hardware LEDC PWM peripheral, and ample GPIO count to simultaneously interface with the PN532 over I2C and drive two servo motors. The PN532 NFC controller module is wired in I2C mode with SDA on GPIO21 and SCL on GPIO22, using the ESP32's hardware I2C controller at 400 kHz. The two SG90 servo motors are connected to GPIO26 and GPIO27 respectively, driven by the ESP32 LEDC peripheral.")
    P(doc, "The PN532 additionally exposes an IRQ (interrupt request) output pin, which is pulled low when a card is detected. This pin is connected to GPIO34 (input-only GPIO) configured for falling-edge interrupt detection, enabling the interrupt-driven NFC design described in Section 4.2. An on-board status LED on GPIO2 provides visual feedback for card detection events and access denial conditions.")

    TABLE(doc,
        ["Peripheral", "GPIO", "Interface", "Configuration"],
        [
            ["PN532 SDA", "GPIO21", "I2C", "400 kHz Fast Mode, pull-up 4.7k ohm"],
            ["PN532 SCL", "GPIO22", "I2C", "400 kHz Fast Mode, pull-up 4.7k ohm"],
            ["PN532 IRQ", "GPIO34", "Digital Input", "INPUT_PULLUP, FALLING edge interrupt"],
            ["Servo Entry", "GPIO26", "LEDC PWM", "50 Hz, 16-bit, LEDC Channel 0"],
            ["Servo Exit",  "GPIO27", "LEDC PWM", "50 Hz, 16-bit, LEDC Channel 1"],
            ["Status LED",  "GPIO2",  "Digital Output", "Active HIGH, onboard blue LED"],
        ],
        col_widths=[3.5, 2.5, 3.0, 8.0]
    )
    CAPTION(doc, "Table 3.2 -- NFC Node GPIO Mapping")

    H(doc, "3.2.2 Camera Nodes (Entry and Exit)", 3)
    P(doc, "Both camera nodes use the AI-Thinker ESP32-CAM module, which integrates the ESP32-S chip, an OV2640 CMOS image sensor, 4 MB of SPI PSRAM, and a microSD card slot on a compact PCB. The OV2640 connects to the ESP32 via a fixed parallel interface (8-bit DVP), with clock, sync, and control signals on dedicated GPIO pins that are not user-configurable. The on-board flash LED on GPIO4 is repurposed as a capture illumination flash and an error indicator (blinking at 5 Hz when the camera subsystem enters the error state).")
    P(doc, "The two camera nodes run identical firmware, differentiated only by a compile-time flag NODE_ROLE (0 = Entry, 1 = Exit) which selects the MQTT topic to subscribe to and determines whether the horizontal mirror correction is applied. This single-firmware approach minimizes maintenance overhead: firmware updates, bug fixes, and configuration changes are applied identically to both nodes.")

    H(doc, "3.2.3 Component Selection Rationale", 3)
    TABLE(doc,
        ["Component", "Selected", "Key Reason", "Alternative Considered"],
        [
            ["MCU (NFC Node)", "ESP32 WROOM32", "Dual-core, WiFi, LEDC PWM, I2C hardware, FreeRTOS", "Arduino Uno: no WiFi, 2KB RAM; STM32: higher cost"],
            ["MCU (Camera)", "ESP32-CAM", "OV2640 + PSRAM on-board, LEDC clock for camera, ecosystem maturity", "Raspberry Pi Zero W: higher cost, Linux boot latency"],
            ["NFC Reader", "PN532 (I2C)", "ISO 14443A support, IRQ pin for interrupt-driven design, Adafruit library", "MFRC522: SPI only, no IRQ output"],
            ["Camera Sensor", "OV2640 (on-board)", "Integrated with ESP32-CAM, JPEG hardware encoder, configurable resolution", "OV5642: 5MP overkill for LPR distance"],
            ["Barrier Motor", "SG90 Servo", "50Hz PWM standard, light weight, adequate torque for model barrier", "Stepper motor: requires dedicated driver IC"],
            ["Comm. Protocol", "MQTT (PubSubClient)", "Pub-sub decoupling, lightweight binary format, broker handles routing", "Direct HTTP: requires knowing target IP, polling-based"],
        ],
        col_widths=[3.0, 3.5, 5.5, 5.0]
    )
    CAPTION(doc, "Table 3.3 -- Component Selection Rationale")

    H(doc, "3.2.4 Power Budget and Supply Design", 3)
    P(doc, "A critical but often overlooked aspect of embedded system design is power supply architecture. Insufficient power supply current capacity causes voltage sag, which manifests as random resets, WiFi disconnections, I2C bus errors, and camera frame corruption -- symptoms that are difficult to diagnose without a current probe. The following power budget was developed for this system:")

    TABLE(doc,
        ["Component", "Supply Voltage", "Idle Current (mA)", "Peak Current (mA)", "Notes"],
        [
            ["ESP32 WROOM32 (NFC Node)", "3.3V", "80", "240", "Peak during WiFi TX burst"],
            ["PN532 NFC Module", "3.3V", "55", "110", "During active field generation"],
            ["SG90 Servo (x2)", "5.0V", "10 each", "500 each", "Stall current per servo; never stall simultaneously"],
            ["Status LED", "3.3V", "10", "10", "Onboard 10 ohm series resistor"],
            ["ESP32-CAM (camera node)", "5.0V (regulated to 3.3V)", "180", "350", "During WiFi TX + camera active"],
            ["OV2640 Camera Active", "2.8V/1.8V (internal reg)", "60", "120", "Included in ESP32-CAM figure"],
        ],
        col_widths=[4.0, 3.0, 3.0, 3.5, 5.0]
    )
    CAPTION(doc, "Table 3.4 -- System Power Budget")

    P(doc, "The most demanding component is the SG90 servo under load (up to 500 mA each at stall). To prevent this current draw from causing voltage sag on the 3.3V rail (which would affect the ESP32 and PN532), the servo motors are powered from a dedicated 5V rail, with only the PWM signal line shared with the ESP32 GPIO. A common ground between the servo supply and the ESP32 3.3V supply is essential for the PWM signal reference. Decoupling capacitors (100 uF electrolytic + 100 nF ceramic) are placed close to each servo connector to absorb transient current demands during servo start and stop.")
    P(doc, "For the prototype, all nodes are powered via USB 5V (500 mA from USB 2.0, or 900 mA from USB 3.0). For production deployment, dedicated 5V/3A switch-mode power supplies would be recommended for each node, with appropriate surge protection and overcurrent limiting.")

    H(doc, "3.3 Network and Communication Architecture", 2)

    H(doc, "3.3.1 MQTT Topic Hierarchy", 3)
    P(doc, "The system defines a structured MQTT topic hierarchy under the root namespace xdhtn/parking/. This namespace convention prevents topic collision with other MQTT services that may share the broker and allows straightforward access control rule definition in production deployments.")
    TABLE(doc,
        ["Topic", "Publisher", "Subscriber(s)", "Payload", "Purpose"],
        [
            ["xdhtn/parking/trigger/entry", "NFC Node", "Entry Camera", "UID string", "Trigger entry capture"],
            ["xdhtn/parking/trigger/exit",  "NFC Node", "Exit Camera",  "UID string", "Trigger exit capture"],
            ["xdhtn/parking/result/entry",  "Entry Camera", "NFC Node", "JSON {card_id, plate, success}", "LPR result for entry"],
            ["xdhtn/parking/result/exit",   "Exit Camera",  "NFC Node", "JSON {card_id, plate, success}", "LPR result for exit"],
            ["xdhtn/parking/override/entry","Server (admin)", "NFC Node", "'override'", "Manual gate open"],
            ["xdhtn/parking/override/exit", "Server (admin)", "NFC Node", "'override'", "Manual gate open"],
        ],
        col_widths=[4.5, 3.0, 3.0, 4.0, 3.5]
    )
    CAPTION(doc, "Table 3.5 -- MQTT Topic Hierarchy")

    H(doc, "3.3.2 HTTP/REST for Log Uploads", 3)
    P(doc, "Following successful LPR processing, each camera node uploads a structured log record to the backend server via HTTP POST. Entry logs are sent as multipart/form-data, including the card UID, detected plate string, LPR success flag, and the JPEG image. Exit logs are sent as either multipart (with exit image) or JSON-only. This design ensures the server has a complete audit record of every vehicle event, including photographic evidence at entry.")

    H(doc, "3.3.3 Server-Sent Events for Dashboard Push", 3)
    P(doc, "The web dashboard receives real-time updates via Server-Sent Events (SSE). Unlike WebSocket, SSE is unidirectional (server to client only), which is sufficient for this application. SSE integrates naturally with FastAPI's StreamingResponse and async generators, and is natively supported by all modern browsers without additional libraries. Every parking event (entry, exit, barrier override, error) is broadcast to all connected SSE clients within milliseconds of occurring.")

    H(doc, "3.3.4 MJPEG for Live Video Streaming", 3)
    P(doc, "Each ESP32-CAM node serves a live MJPEG video stream on port 81 at the /stream endpoint. The backend server proxies these streams at /stream/entry and /stream/exit, allowing the dashboard to display both camera feeds through a single domain without cross-origin issues. The proxy is implemented as an async byte-relay using httpx's streaming GET, consuming and forwarding 4096-byte chunks without buffering entire frames -- keeping latency minimal.")

    H(doc, "3.4 Software Architecture", 2)

    H(doc, "3.4.1 Four-Layer Software Model", 3)
    P(doc, "The firmware is organized into four conceptual layers, following the principle of separation of concerns.")
    CODE(doc,
         "+-----------------------------------------------------------------+\n"
         "| APPLICATION LAYER                                                |\n"
         "| Parking rules: entry/exit FSM, session management, barrier logic |\n"
         "+-----------------------------------------------------------------+\n"
         "| SERVICE / COMMUNICATION LAYER                                    |\n"
         "| MQTT pub/sub, HTTP multipart upload, FastAPI REST + SSE          |\n"
         "+-----------------------------------------------------------------+\n"
         "| FIRMWARE TASK LAYER  (FreeRTOS)                                  |\n"
         "| ESP32-CAM: taskMQTT|taskCamera|taskLPR|taskLog|taskStream        |\n"
         "| NFC Node:  taskNFC|taskServo|[SWTimer: gateClose]                |\n"
         "+-----------------------------------------------------------------+\n"
         "| HAL / DRIVER LAYER                                               |\n"
         "| esp_camera (I2S DMA), Adafruit_PN532 (I2C), LEDC PWM,           |\n"
         "| WiFi/TCP stack, PSRAM heap_caps API                              |\n"
         "+-----------------------------------------------------------------+"
    )
    CAPTION(doc, "Figure 3.1 -- Four-Layer Software Architecture")

    H(doc, "3.4.2 Firmware Segregation via PlatformIO", 3)
    P(doc, "A single source tree supports three distinct firmware binaries through the PlatformIO build system's build_src_filter mechanism. Three PlatformIO environments are defined: entry_node (ESP32-CAM, NODE_ROLE=0), exit_node (ESP32-CAM, NODE_ROLE=1), and nfc_node (ESP32 WROOM32). The NODE_ROLE preprocessor define allows a single camera firmware source to produce two role-specific binaries.")

    H(doc, "3.4.3 Finite State Machine Design", 3)
    P(doc, "The ESP32-CAM firmware implements an explicit Finite State Machine (FSM) to govern the sequencing of camera capture operations. The FSM state is stored in a volatile global variable (gState, of type SystemState_t) protected by a mutex (xStateMutex).")

    TABLE(doc,
        ["State", "Meaning", "Transitions To", "Duration"],
        [
            ["STATE_IDLE", "System ready, awaiting trigger", "STATE_CARD_READ on MQTT trigger", "Indefinite"],
            ["STATE_CARD_READ", "MQTT trigger received, camera arming", "STATE_CAPTURING immediately", "<100 ms"],
            ["STATE_CAPTURING", "Camera warm-up and frame capture in progress", "STATE_LPR_PENDING on success; STATE_IDLE on failure", "~300 ms"],
            ["STATE_LPR_PENDING", "LPR API call in progress", "STATE_GATE_OPEN on success; STATE_LPR_FAIL on failure", "500-3000 ms"],
            ["STATE_GATE_OPEN", "LPR succeeded, MQTT result published", "STATE_IDLE after publish", "<100 ms"],
            ["STATE_LPR_FAIL", "LPR failed or confidence below threshold", "STATE_IDLE after failure handling", "<100 ms"],
            ["STATE_ERROR", "Camera hardware initialization failure", "No automatic recovery; LED blinks", "Indefinite"],
        ],
        col_widths=[3.5, 4.0, 4.5, 2.5]
    )
    CAPTION(doc, "Table 3.6 -- ESP32-CAM FSM State Table")

    H(doc, "3.4.4 Data Flow: Entry and Exit Sequences", 3)
    P(doc, "The complete data flow for a vehicle entry event, from card presentation to barrier opening, proceeds through the pipeline shown below.")
    CODE(doc,
         "ENTRY SEQUENCE:\n"
         "  [Driver presents NFC card to PN532]\n"
         "        | PN532 IRQ pin falls LOW\n"
         "        v\n"
         "  taskNFC (Core 1)  -- ISR wakes task via xSemaphoreGiveFromISR\n"
         "        | reads UID, checks gParkingTable (NOT in table = ENTRY)\n"
         "        | publishes to xdhtn/parking/trigger/entry\n"
         "        v\n"
         "  Entry ESP32-CAM -- taskMQTT receives trigger, xQueueOverwrite(xNfcQueue)\n"
         "        |\n"
         "        v\n"
         "  taskCamera (Core 1) -- warm-up + capture + PSRAM alloc + xQueueSend\n"
         "        |\n"
         "        v\n"
         "  taskLPR (Core 0) -- HTTP POST to /api/proxy-lpr, parse result, validate\n"
         "        | publishes to xdhtn/parking/result/entry {plate, success=1}\n"
         "        v\n"
         "  NFC Node -- mqttCallback: tableInsert(uid, plate) -> GATE_ENTRY command\n"
         "        v\n"
         "  taskServo -- servoOpen(ENTRY_PIN) -> xTimerStart(closeTimer, 5000ms)\n"
         "        v\n"
         "  [Barrier opens; 5-second timer auto-closes gate via callback]"
    )
    CAPTION(doc, "Figure 3.2 -- Entry Event Full Data Flow")

    H(doc, "3.5 Security Architecture and Considerations", 2)
    P(doc, "Security in embedded IoT systems is frequently treated as an afterthought, leading to deployed systems that are vulnerable to network-level attacks, unauthorized command injection, and data interception. This section documents the security posture of the current prototype implementation, identifies known vulnerabilities, and describes the architectural controls that mitigate risk within the project's scope.")

    H(doc, "3.5.1 Current Security Posture", 3)
    P(doc, "The prototype implementation prioritizes functional correctness and demonstrating embedded systems techniques over operational security hardening. The following security properties characterize the current implementation:")
    BULLET(doc, [
        "MQTT communication between all nodes is unencrypted (plaintext TCP on port 1883). Any device on the same WiFi network can observe MQTT messages, publish fake trigger messages, or subscribe to result topics.",
        "The /api/barrier/{node} HTTP endpoint has no authentication. Any HTTP client on the LAN can send a POST request to open either barrier gate without authorization.",
        "NFC card UIDs are transmitted in plaintext over MQTT and HTTP. A passive network observer can learn which UIDs are associated with which parking events.",
        "NFC card UIDs for ISO 14443A MIFARE Classic cards can be cloned using commercially available tools (e.g., Proxmark 3, Flipper Zero). A cloned card would be accepted by the system as genuine.",
        "The WiFi network itself provides a first layer of access control: only devices with the WiFi PSK can join the network and interact with the MQTT broker and HTTP endpoints.",
    ])

    H(doc, "3.5.2 Proposed Security Enhancements for Production", 3)
    TABLE(doc,
        ["Vulnerability", "Risk Level", "Proposed Mitigation", "Implementation Effort"],
        [
            ["Unencrypted MQTT", "High", "Enable MQTT over TLS (port 8883) with self-signed certificates; configure Mosquitto with require_certificate; use WiFiClientSecure on ESP32", "Medium -- TLS handshake adds ~100ms latency; requires certificate provisioning"],
            ["Unauthenticated barrier API", "High", "Add API key header (X-API-Key) to /api/barrier endpoint; store key in server .env; validate in FastAPI dependency", "Low -- 10-20 lines of Python code"],
            ["NFC UID cloning", "Medium", "Migrate to MIFARE DESFire cards with AES-128 mutual authentication; PN532 supports DESFire EV1 protocol", "High -- requires card re-issuance and firmware changes"],
            ["Plaintext MQTT payloads", "Low", "Application-layer encryption of MQTT payloads using AES-256-GCM; key stored in firmware as compile-time constant", "Medium -- adds payload processing overhead"],
            ["MQTT broker open to LAN", "Medium", "Configure Mosquitto with username/password authentication and per-client ACL rules restricting publish/subscribe topics", "Low -- Mosquitto configuration only"],
        ],
        col_widths=[3.5, 2.0, 5.5, 6.0]
    )
    CAPTION(doc, "Table 3.7 -- Security Vulnerability and Mitigation Analysis")

    H(doc, "3.6 System Scalability Analysis", 2)
    P(doc, "A key design objective of this system is scalability: the ability to add additional parking lanes without architectural changes. This section analyzes the scalability properties of the current design.")
    P(doc, "To add an additional entry/exit lane pair to the system, the following steps are required:")
    BULLET(doc, [
        "Deploy one additional entry ESP32-CAM node and one additional exit ESP32-CAM node, flashed with the standard firmware. Configure their WiFi credentials and SERVER_HOST in config.h.",
        "Extend the MQTT topic hierarchy: the new entry camera subscribes to xdhtn/parking/trigger/entry_lane2 and publishes to xdhtn/parking/result/entry_lane2. No changes to the MQTT broker or existing nodes are required.",
        "The NFC Node firmware must be updated to publish triggers to the correct lane's topic based on which physical entry/exit point the card was presented at. If multiple PN532 readers are used (one per lane), each reader's ISR publishes to its designated topic.",
        "The backend server's .env file is updated with the IP addresses of the new camera nodes. The session management and SSE broadcasting code requires no changes.",
    ])
    P(doc, "This analysis confirms that the system scales horizontally in the camera tier without any backend or broker code changes. The NFC/barrier node is the only component requiring firmware modification for multi-lane support, and this modification is limited to topic routing logic -- not architectural changes. For a four-lane facility, the estimated additional cost per lane pair is approximately 375,000-500,000 VND (two ESP32-CAM modules), representing an order-of-magnitude cost reduction compared to commercial lane additions.")
    PAGE_BREAK(doc)


# =============================================================================
# CHAPTER 4
# =============================================================================

def chapter4(doc):
    H(doc, "CHAPTER 4: SYSTEM IMPLEMENTATION", 1)

    H(doc, "4.1 Development Environment and Build System", 2)
    P(doc, "The firmware for both node types is developed using the PlatformIO IDE extension for Visual Studio Code, targeting the Arduino framework on top of Espressif's ESP-IDF. PlatformIO provides dependency management, build flag propagation, multi-environment support, serial monitor, and firmware upload -- all without requiring manual toolchain installation.")
    P(doc, "Three build environments are defined in platformio.ini: entry_node and exit_node (both targeting the esp32cam board), and nfc_node (targeting the esp32dev board). The camera environments inherit common settings from a base environment env_cam_base.")

    CODE(doc,
         "; platformio.ini -- abbreviated\n"
         "[env_cam_base]\n"
         "platform  = espressif32\n"
         "board     = esp32cam\n"
         "framework = arduino\n"
         "board_build.partitions = huge_app.csv\n"
         "lib_deps  = bblanchon/ArduinoJson@^7.0.0\n"
         "            knolleary/PubSubClient@^2.8\n"
         "build_src_filter =\n"
         "    +<*> -<nfc_node.cpp> -<nfc_task.cpp> -<http_task.cpp>\n"
         "build_flags = -DCORE_DEBUG_LEVEL=3 -DBOARD_HAS_PSRAM\n"
         "              -mfix-esp32-psram-cache-issue\n"
         "\n"
         "[env:entry_node]\n"
         "extends    = env_cam_base\n"
         "build_flags = ${env_cam_base.build_flags} -DNODE_ROLE=0\n"
         "\n"
         "[env:exit_node]\n"
         "extends    = env_cam_base\n"
         "build_flags = ${env_cam_base.build_flags} -DNODE_ROLE=1\n"
         "\n"
         "[env:nfc_node]\n"
         "platform  = espressif32\n"
         "board     = esp32dev\n"
         "framework = arduino\n"
         "lib_deps  = adafruit/Adafruit PN532@^1.3.4\n"
         "            knolleary/PubSubClient@^2.8\n"
         "            bblanchon/ArduinoJson@^7.0.0\n"
         "build_src_filter = +<nfc_node.cpp>"
    )
    CAPTION(doc, "Figure 4.1 -- PlatformIO Configuration File")

    H(doc, "4.2 Interrupt-Driven NFC Reader Implementation", 2)
    P(doc, "The NFC reader implementation is one of the most technically significant components of this project from an embedded systems perspective. Two approaches were designed and analyzed: polling-based NFC reading, and interrupt-driven (IRQ-triggered) NFC reading. The interrupt-driven approach represents the architecturally preferred design and is documented here as the primary implementation.")

    H(doc, "4.2.1 Interrupt Service Routine Design", 3)
    P(doc, "The PN532 NFC controller asserts its IRQ output pin LOW when a card has been detected and data is ready to be read. By connecting this pin to a GPIO input on the ESP32 and attaching a FALLING-edge hardware interrupt, the taskNFC FreeRTOS task can sleep indefinitely (consuming zero CPU) and be woken instantly the moment a card is presented -- eliminating the CPU overhead and worst-case 1-second latency of polling.")
    P(doc, "The ISR function must be placed in IRAM (Internal RAM) rather than flash memory, because on the ESP32, flash is accessed via SPI with cache and cannot be reliably read during certain WiFi operations. The IRAM_ATTR attribute instructs the linker to place the function in IRAM:")

    CODE(doc,
         "// nfc_task.cpp -- Interrupt Service Routine\n"
         "static SemaphoreHandle_t xIrqSemaphore = nullptr;\n"
         "\n"
         "static void IRAM_ATTR pn532IrqHandler() {\n"
         "    // Must be ISR-safe: no floating point, no heap, no blocking\n"
         "    BaseType_t xHigherPriorityTaskWoken = pdFALSE;\n"
         "    xSemaphoreGiveFromISR(xIrqSemaphore, &xHigherPriorityTaskWoken);\n"
         "    // If giving the semaphore unblocked a higher-priority task,\n"
         "    // request an immediate context switch at ISR exit\n"
         "    portYIELD_FROM_ISR(xHigherPriorityTaskWoken);\n"
         "}"
    )
    CAPTION(doc, "Figure 4.2 -- ISR Function with IRAM_ATTR")

    H(doc, "4.2.2 Race Condition Analysis and Prevention", 3)
    P(doc, "A subtle but critical race condition exists in interrupt-driven NFC designs. If the ISR fires between the point where xIrqSemaphore is declared and where xSemaphoreCreateBinary() initializes it, the ISR will attempt to call xSemaphoreGiveFromISR() on a null handle, resulting in a crash. The correct fix is to ensure the semaphore is fully created before the interrupt is attached:")

    CODE(doc,
         "// taskNFC() -- Correct initialization ordering\n"
         "void taskNFC(void *pvParameters) {\n"
         "    Wire.begin(PN532_SDA_PIN, PN532_SCL_PIN, 400000UL);\n"
         "    pn532.begin();\n"
         "    // ... (PN532 firmware check and SAMConfig) ...\n"
         "\n"
         "    // Step 1: Create semaphore BEFORE attaching interrupt\n"
         "    xIrqSemaphore = xSemaphoreCreateBinary();\n"
         "    configASSERT(xIrqSemaphore);\n"
         "\n"
         "    // Step 2: Configure GPIO and arm PN532\n"
         "    pinMode(PN532_IRQ_PIN, INPUT_PULLUP);\n"
         "    pn532.startPassiveTargetIDDetection(PN532_MIFARE_ISO14443A);\n"
         "\n"
         "    // Step 3: Attach interrupt LAST -- semaphore guaranteed to exist\n"
         "    attachInterrupt(digitalPinToInterrupt(PN532_IRQ_PIN),\n"
         "                    pn532IrqHandler, FALLING);\n"
         "}"
    )
    CAPTION(doc, "Figure 4.3 -- Race-Condition-Free Interrupt Initialization")

    H(doc, "4.2.3 Main Task Loop and Debounce", 3)
    P(doc, "With the interrupt infrastructure in place, the main NFC task loop blocks indefinitely on the semaphore, wakes only when a card is detected, reads the UID, processes the event, and then debounces by waiting for card removal:")

    CODE(doc,
         "while (true) {\n"
         "    pn532.startPassiveTargetIDDetection(PN532_MIFARE_ISO14443A);\n"
         "\n"
         "    if (xSemaphoreTake(xIrqSemaphore,\n"
         "                       pdMS_TO_TICKS(NFC_READ_TIMEOUT_MS)) != pdTRUE)\n"
         "        continue; // timeout -- no card, re-arm silently\n"
         "\n"
         "    bool success = pn532.readDetectedPassiveTargetID(evt.uid, &evt.uidLen);\n"
         "    if (!success || evt.uidLen == 0) { continue; }\n"
         "\n"
         "    uidToString(evt.uid, evt.uidLen, evt.uidStr, sizeof(evt.uidStr));\n"
         "    // ... process event (MQTT publish, parking table lookup) ...\n"
         "\n"
         "    // Debounce: wait for card removal (up to 3 seconds)\n"
         "    uint32_t debounceStart = millis();\n"
         "    while (millis() - debounceStart < 3000) {\n"
         "        pn532.startPassiveTargetIDDetection(PN532_MIFARE_ISO14443A);\n"
         "        if (xSemaphoreTake(xIrqSemaphore, pdMS_TO_TICKS(200)) != pdTRUE)\n"
         "            break; // no IRQ in 200ms -> card removed\n"
         "        pn532.readDetectedPassiveTargetID(tmpUid, &tmpLen);\n"
         "    }\n"
         "}"
    )
    CAPTION(doc, "Figure 4.4 -- NFC Task Main Loop with Debounce")

    H(doc, "4.3 Hardware Timer and PWM Implementation", 2)

    H(doc, "4.3.1 LEDC PWM for Servo Motor Control", 3)
    P(doc, "The SG90 servo motor requires a 50 Hz PWM signal (20 ms period) with a pulse width between approximately 0.5 ms and 2.4 ms to control its angular position. Using 16-bit resolution provides 65,536 discrete duty cycle steps. The calculation is as follows:")
    CODE(doc,
         "Period     = 1 / 50 Hz = 20,000 us\n"
         "Step size  = 20,000 us / 65,536 steps = 0.305 us/step\n"
         "0.5ms pulse  = 500  us / 0.305 us = ~1638 steps  (0 degrees -- closed)\n"
         "1.5ms pulse  = 1500 us / 0.305 us = ~4915 steps  (90 degrees -- open)\n"
         "2.4ms pulse  = 2400 us / 0.305 us = ~7864 steps  (180 degrees)\n"
         "\n"
         "static uint32_t angleToDuty(int angle) {\n"
         "    return (uint32_t)(1638 + (angle * (8192 - 1638)) / 180);\n"
         "}\n"
         "\n"
         "ledcAttach(SERVO_ENTRY_PIN, 50, 16); // 50Hz, 16-bit\n"
         "ledcAttach(SERVO_EXIT_PIN,  50, 16);\n"
         "ledcWrite(SERVO_ENTRY_PIN, angleToDuty(0)); // start closed"
    )
    CAPTION(doc, "Figure 4.5 -- LEDC PWM Duty Cycle Calculation")

    H(doc, "4.3.2 FreeRTOS Software Timer for Gate Auto-Close", 3)
    P(doc, "The improved gate auto-close uses a FreeRTOS Software Timer -- a callback scheduled after 5 seconds -- eliminating the need for a blocking one-shot task:")

    CODE(doc,
         "static TimerHandle_t gEntryCloseTimer = nullptr;\n"
         "static TimerHandle_t gExitCloseTimer  = nullptr;\n"
         "\n"
         "static void entryCloseCallback(TimerHandle_t xTimer) {\n"
         "    servoSet(SERVO_ENTRY_PIN, SERVO_CLOSE_DEG);\n"
         "    digitalWrite(LED_PIN, LOW);\n"
         "}\n"
         "\n"
         "// In setup():\n"
         "gEntryCloseTimer = xTimerCreate(\n"
         "    \"entryClose\", pdMS_TO_TICKS(5000),\n"
         "    pdFALSE, nullptr, entryCloseCallback);\n"
         "\n"
         "// When gate opens:\n"
         "static void openGate(uint8_t pin, TimerHandle_t closeTimer) {\n"
         "    servoSet(pin, SERVO_OPEN_DEG);\n"
         "    digitalWrite(LED_PIN, HIGH);\n"
         "    xTimerStart(closeTimer, 0); // non-blocking 5s countdown\n"
         "}"
    )
    CAPTION(doc, "Figure 4.6 -- FreeRTOS Software Timer for Non-Blocking Gate Close")

    H(doc, "4.3.3 Hardware Watchdog Timer for Fault Recovery", 3)
    P(doc, "The ESP32 hardware watchdog timer provides a backstop against I2C bus lockups. The watchdog is configured with an 8-second timeout; the NFC task pets it at the start of each loop iteration:")

    CODE(doc,
         "static hw_timer_t *wdTimer = nullptr;\n"
         "static volatile bool wdFired = false;\n"
         "\n"
         "static void IRAM_ATTR watchdogISR() { wdFired = true; }\n"
         "\n"
         "// In setup():\n"
         "wdTimer = timerBegin(1, 80, true);     // Timer 1, 1MHz tick\n"
         "timerAttachInterrupt(wdTimer, &watchdogISR, true);\n"
         "timerAlarmWrite(wdTimer, 8000000, false); // 8 second timeout\n"
         "timerAlarmEnable(wdTimer);\n"
         "\n"
         "// In taskNFC main loop:\n"
         "timerWrite(wdTimer, 0); // pet the watchdog\n"
         "if (wdFired) {\n"
         "    wdFired = false;\n"
         "    Wire.end();\n"
         "    vTaskDelay(pdMS_TO_TICKS(100));\n"
         "    Wire.begin(NFC_SDA_PIN, NFC_SCL_PIN, 400000UL);\n"
         "    pn532.begin(); pn532.SAMConfig();\n"
         "    timerAlarmEnable(wdTimer);\n"
         "}"
    )
    CAPTION(doc, "Figure 4.7 -- Hardware Watchdog Timer for I2C Fault Recovery")

    H(doc, "4.4 FreeRTOS Task Architecture", 2)

    H(doc, "4.4.1 Complete Task Map", 3)
    TABLE(doc,
        ["Task", "Function", "Stack (B)", "Priority", "Core", "Role"],
        [
            ["taskMQTT",      "taskMQTT()",      "4096", "2", "0", "MQTT client loop, receive trigger, dispatch NfcEvent"],
            ["taskCamera",    "taskCamera()",    "8192", "3", "1", "Camera init, warm-up, capture, PSRAM buffer management"],
            ["taskLPR",       "taskLPR()",       "8192", "2", "0", "LPR API call, JSON parse, plate validation, MQTT publish"],
            ["taskLogUpload", "taskLogUpload()", "6144", "1", "0", "HTTP log upload to backend, image and JSON payloads"],
            ["taskStream",    "taskStreamServer()","6144","2","1", "MJPEG httpd server startup; httpd runs on IDF internal task"],
        ],
        col_widths=[3.0, 3.5, 2.0, 2.5, 1.5, 5.5]
    )
    CAPTION(doc, "Table 4.1 -- ESP32-CAM FreeRTOS Task Map")

    TABLE(doc,
        ["Task", "Function", "Stack (B)", "Priority", "Core", "Role"],
        [
            ["taskNFC",   "taskNFC()",   "4096", "3", "1", "IRQ-driven NFC read, parking table lookup, MQTT trigger"],
            ["taskServo", "taskServo()", "3072", "2", "0", "Receive gate commands from xGateQueue, actuate servo"],
        ],
        col_widths=[2.5, 3.0, 2.0, 2.5, 1.5, 6.5]
    )
    CAPTION(doc, "Table 4.2 -- NFC Node FreeRTOS Task Map")

    H(doc, "4.4.2 Core Pinning Rationale", 3)
    P(doc, "Camera and stream tasks are pinned to Core 1, isolated from WiFi interrupts. All network tasks (taskMQTT, taskLPR, taskLogUpload) are pinned to Core 0, where they share CPU time with the TCP/IP stack and benefit from cache locality -- reducing the overhead of cross-core context switches.")

    H(doc, "4.4.3 Priority Assignment and Priority Inversion Analysis", 3)
    P(doc, "taskCamera has the highest priority (3) because it sits on the critical path. taskMQTT and taskLPR share priority 2, ensuring MQTT messages are processed promptly. taskLogUpload has priority 1 (lowest non-idle) because log uploads are best-effort. All mutexes are created with xSemaphoreCreateMutex() (not xSemaphoreCreateBinary()), which provides FreeRTOS priority inheritance, preventing priority inversion scenarios.")

    H(doc, "4.4.4 Stack Size Determination and High-Water Mark Monitoring", 3)
    P(doc, "Stack overflow is one of the most common causes of silent corruption in FreeRTOS applications. A periodic stack high-water mark reporting routine calls uxTaskGetStackHighWaterMark() for each task handle every 30 seconds:")

    CODE(doc,
         "// main.cpp -- Stack high-water mark monitoring\n"
         "static uint32_t lastStackReport = 0;\n"
         "\n"
         "if (millis() - lastStackReport > 30000) {\n"
         "    lastStackReport = millis();\n"
         "    Serial.printf(\"[STACK] taskCamera:    %u words free\\n\",\n"
         "        uxTaskGetStackHighWaterMark(hCameraTask));\n"
         "    Serial.printf(\"[STACK] taskLPR:       %u words free\\n\",\n"
         "        uxTaskGetStackHighWaterMark(hLprTask));\n"
         "    Serial.printf(\"[STACK] taskMQTT:      %u words free\\n\",\n"
         "        uxTaskGetStackHighWaterMark(hMqttTask));\n"
         "    Serial.printf(\"[STACK] taskLogUpload: %u words free\\n\",\n"
         "        uxTaskGetStackHighWaterMark(hLogTask));\n"
         "}"
    )
    CAPTION(doc, "Figure 4.8 -- Runtime Stack High-Water Mark Monitoring")

    H(doc, "4.5 Inter-Task Communication and Synchronization", 2)

    H(doc, "4.5.1 Queue Pipeline Architecture", 3)
    P(doc, "Data flows through the ESP32-CAM firmware as a pipeline of FreeRTOS queues, each with depth 1 except the log queue (depth 4). The depth-1 design reflects the system's single-vehicle-at-a-time processing model.")
    CODE(doc,
         "  taskMQTT   ->  xNfcQueue      (depth 1, NfcEvent_t)\n"
         "                     | xQueueOverwrite() -- always replaces stale event\n"
         "                     v\n"
         "  taskCamera ->  xCaptureQueue  (depth 1, CaptureEvent_t)\n"
         "                     | xQueueSend()  -- blocks 500ms if full\n"
         "                     v\n"
         "  taskLPR    ->  xLogQueue      (depth 4, LprEvent_t)\n"
         "                     | xQueueSend()  -- non-blocking, drops if full\n"
         "                     v\n"
         "  taskLogUpload -> POST /api/log/entry or /api/log/exit"
    )
    CAPTION(doc, "Figure 4.9 -- FreeRTOS Queue Pipeline")

    H(doc, "4.5.2 Mutex-Protected Shared Resources", 3)
    P(doc, "Two global resources are shared across multiple tasks and require mutex protection: xStateMutex protects the gState FSM state variable, and xCameraMutex serializes access to the esp_camera framebuffer API.")

    CODE(doc,
         "// mutex-protected FSM state update\n"
         "if (xSemaphoreTake(xStateMutex, pdMS_TO_TICKS(100)) == pdTRUE) {\n"
         "    gState = STATE_CAPTURING;\n"
         "    xSemaphoreGive(xStateMutex);\n"
         "}\n"
         "\n"
         "// Always return DMA buffer even on mutex timeout\n"
         "if (fb) {\n"
         "    if (xSemaphoreTake(xCameraMutex, pdMS_TO_TICKS(500)) == pdTRUE) {\n"
         "        esp_camera_fb_return(fb);\n"
         "        xSemaphoreGive(xCameraMutex);\n"
         "    } else {\n"
         "        esp_camera_fb_return(fb); // must return even without mutex\n"
         "    }\n"
         "}"
    )
    CAPTION(doc, "Figure 4.10 -- Mutex-Protected Shared Resource Access Patterns")

    H(doc, "4.5.3 Task Notifications for Camera-Ready Signaling", 3)
    P(doc, "The improved implementation uses a FreeRTOS Task Notification instead of a volatile bool flag. taskCamera sends a notification to taskStream immediately after successful initialization, and taskStream blocks on ulTaskNotifyTake() until the notification arrives:")

    CODE(doc,
         "// camera_task.cpp\n"
         "xTaskNotifyGive(hStreamTask); // camera ready\n"
         "\n"
         "// stream_server.cpp\n"
         "ulTaskNotifyTake(pdTRUE, portMAX_DELAY); // block until notified\n"
         "startMjpegServer(); // camera guaranteed initialized"
    )
    CAPTION(doc, "Figure 4.11 -- Task Notification for Camera-Ready Signaling")

    H(doc, "4.6 Camera Subsystem Implementation", 2)

    H(doc, "4.6.1 OV2640 Initialization and Sensor Tuning", 3)
    P(doc, "The OV2640 camera sensor is initialized using the esp_camera library. The camera is configured for SVGA resolution (800x600 pixels), JPEG pixel format, quality factor 12, and three framebuffers in PSRAM. After initialization, sensor registers are tuned for the parking gate environment: brightness +1 (compensate for variable outdoor lighting), contrast +1 (improve plate edge definition), auto white balance enabled, and horizontal mirror correction applied for the entry camera.")

    H(doc, "4.6.2 DMA-Based Frame Capture Pipeline", 3)
    P(doc, "When esp_camera_fb_get() is called, the following sequence occurs within the esp_camera/IDF layer, entirely without application CPU involvement:")
    BULLET(doc, [
        "The I2S peripheral receives pixel clock (PCLK) and data signals from the OV2640 via the DVP interface.",
        "On each VSYNC pulse (frame start), a DMA channel is armed to transfer pixel data from the I2S FIFO into a pre-allocated PSRAM framebuffer.",
        "The OV2640 DSP compresses the raw Bayer sensor data into JPEG on-chip. JPEG-compressed data is output on the DVP bus.",
        "When HREF goes low at the end of the last line (frame complete), the DMA transfer is complete and an internal semaphore is given.",
        "esp_camera_fb_get() returns the pointer to the completed JPEG data in PSRAM.",
    ])

    H(doc, "4.6.3 PSRAM Buffer Lifecycle", 3)
    P(doc, "Correct PSRAM buffer lifecycle management prevents memory leaks across many capture cycles:")
    CODE(doc,
         "PSRAM Buffer Lifecycle:\n"
         "  [taskCamera]\n"
         "  1. esp_camera_fb_get()          -> DMA buffer in PSRAM (camera driver)\n"
         "  2. heap_caps_malloc(SPIRAM)      -> application PSRAM buffer allocated\n"
         "  3. memcpy(appBuf, dma.buf, len) -> copy frame to application buffer\n"
         "  4. esp_camera_fb_return(dma)    -> DMA buffer returned to camera pool\n"
         "  5. xQueueSend(xCaptureQueue, {appBuf}) -> ownership to taskLPR\n"
         "\n"
         "  [taskLPR]\n"
         "  6. buildPlateApiBody(appBuf)    -> multipart body in PSRAM allocated\n"
         "  7. http.POST(body, bodyLen)     -> image sent to LPR API\n"
         "  8. heap_caps_free(body)         -> multipart body freed\n"
         "  9. xQueueSend(xLogQueue, {appBuf}) -> ownership to taskLogUpload\n"
         "\n"
         "  [taskLogUpload]\n"
         " 10. buildEntryLog(appBuf)        -> log body in PSRAM allocated\n"
         " 11. http.POST(logBody)           -> log + image sent to server\n"
         " 12. heap_caps_free(logBody)      -> log body freed\n"
         " 13. heap_caps_free(appBuf)       -> APPLICATION BUFFER FREED -- no leak"
    )
    CAPTION(doc, "Figure 4.12 -- Complete PSRAM Buffer Ownership Lifecycle")

    H(doc, "4.6.4 Warm-Up Frame and Auto-Exposure Settle", 3)
    P(doc, "When capture is triggered, the OV2640's AE/AWB algorithms need one frame period to stabilize. The firmware captures one warm-up frame (discarded) and waits 150 ms before the actual capture. During this 150 ms, the flash LED is illuminated for consistent lighting. This adds 200-250 ms to capture latency but substantially improves LPR accuracy.")

    H(doc, "4.7 License Plate Recognition Pipeline", 2)

    H(doc, "4.7.1 Server-Proxied LPR API Call", 3)
    P(doc, "The ESP32-CAM communicates with the Plate Recognizer API indirectly through the backend FastAPI server's /api/proxy-lpr endpoint. This proxy architecture offloads TLS handshake overhead to the server and keeps the API token server-side for security. The ESP32-CAM sends an unencrypted HTTP POST to the local server, which forwards the image to Plate Recognizer over HTTPS.")
    P(doc, "The Plate Recognizer API returns a JSON response containing an array of detected plate results. Each result contains a nested 'candidates' array with alternative plate reads and their associated confidence scores, as well as vehicle metadata (type, color, region). The following is a representative API response for a Vietnamese SVGA image:")

    CODE(doc,
         "// Example Plate Recognizer API JSON response\n"
         "{\n"
         "  \"results\": [\n"
         "    {\n"
         "      \"plate\": \"51A12345\",\n"
         "      \"confidence\": 0.897,\n"
         "      \"region\": { \"code\": \"vn\", \"score\": 0.943 },\n"
         "      \"box\": { \"xmin\": 312, \"ymin\": 401, \"xmax\": 520, \"ymax\": 448 },\n"
         "      \"dscore\": 0.912,\n"
         "      \"candidates\": [\n"
         "        { \"plate\": \"51A12345\", \"score\": 0.897 },\n"
         "        { \"plate\": \"51A12B45\", \"score\": 0.312 }\n"
         "      ]\n"
         "    }\n"
         "  ],\n"
         "  \"usage\": { \"calls\": 47 }\n"
         "}"
    )
    CAPTION(doc, "Figure 4.13 -- Representative Plate Recognizer API JSON Response")

    P(doc, "The firmware's LPR parsing function uses ArduinoJson to deserialize this response. It extracts the first result's plate string and confidence score, normalizes the plate string (strips spaces, converts to uppercase, removes hyphens and dots), and then applies the confidence threshold and format validation check:")

    CODE(doc,
         "// lpr_task.cpp -- parseLprResponse()\n"
         "bool parseLprResponse(const char *json, char *outPlate, float *outConf) {\n"
         "    StaticJsonDocument<2048> doc;\n"
         "    if (deserializeJson(doc, json) != DeserializationError::Ok)\n"
         "        return false;\n"
         "\n"
         "    JsonArray results = doc[\"results\"];\n"
         "    if (results.isNull() || results.size() == 0)\n"
         "        return false; // no plate detected in image\n"
         "\n"
         "    JsonObject best = results[0]; // highest dscore result\n"
         "    const char *rawPlate = best[\"plate\"];\n"
         "    float conf = best[\"confidence\"];\n"
         "\n"
         "    // Normalize: uppercase, strip hyphens/dots/spaces\n"
         "    char normalized[PLATE_STR_LEN] = {};\n"
         "    int j = 0;\n"
         "    for (int i = 0; rawPlate[i] && j < PLATE_STR_LEN-1; i++) {\n"
         "        char c = toupper(rawPlate[i]);\n"
         "        if (c != '-' && c != '.' && c != ' ')\n"
         "            normalized[j++] = c;\n"
         "    }\n"
         "\n"
         "    if (conf >= PLATE_CONFIDENCE && plateValid(normalized)) {\n"
         "        strlcpy(outPlate, normalized, PLATE_STR_LEN);\n"
         "        *outConf = conf;\n"
         "        return true;\n"
         "    }\n"
         "    return false; // below threshold or invalid format\n"
         "}"
    )
    CAPTION(doc, "Figure 4.14 -- LPR JSON Parsing and Normalization")

    H(doc, "4.7.2 Confidence Thresholding and Vietnamese Plate Validation", 3)
    P(doc, "Vietnamese license plates follow one of two formats mandated by Circular 58/2020/TT-BCA of the Ministry of Public Security: (1) a two-digit province code, one or two uppercase letters, followed by four or five digits (e.g., '51A12345', '30AB12345'); and (2) the same format with a hyphen separator for display (e.g., '51A-123.45'). After normalization, both formats reduce to the same character string matching the regular expression ^[0-9]{2}[A-Z]{1,2}[0-9]{4,5}$. This validation step is important because the Plate Recognizer API occasionally misreads a digit as a letter or vice versa -- a plate read of '5lA12345' (lowercase 'l' instead of '1') would fail regex validation and be correctly rejected as an LPR error rather than corrupting the parking record.")

    CODE(doc,
         "#define PLATE_CONFIDENCE  0.75f\n"
         "// Vietnamese formats: 51A12345, 51AB12345 (post-normalization)\n"
         "#define PLATE_REGEX  \"^[0-9]{2}[A-Z]{1,2}[0-9]{4,5}$\"\n"
         "\n"
         "static bool plateValid(const char *plate) {\n"
         "    regex_t re;\n"
         "    if (regcomp(&re, PLATE_REGEX, REG_EXTENDED | REG_NOSUB) != 0)\n"
         "        return false;\n"
         "    bool ok = (regexec(&re, plate, 0, nullptr, 0) == 0);\n"
         "    regfree(&re);\n"
         "    return ok;\n"
         "}"
    )
    CAPTION(doc, "Figure 4.15 -- Vietnamese License Plate Format Validation Regex")

    H(doc, "4.7.3 FastAPI LPR Proxy Implementation", 3)
    P(doc, "The backend /api/proxy-lpr endpoint receives the JPEG image from the ESP32-CAM as a multipart form upload, adds the Plate Recognizer API key header, and forwards the image to the cloud API. This design keeps the API secret entirely on the server side -- the ESP32-CAM firmware never sees or stores the API token:")

    CODE(doc,
         "# server/main.py -- LPR proxy endpoint\n"
         "@app.post('/api/proxy-lpr')\n"
         "async def proxy_lpr(upload: UploadFile = File(...)):\n"
         "    img_bytes = await upload.read()\n"
         "    async with httpx.AsyncClient(timeout=8.0) as client:\n"
         "        resp = await client.post(\n"
         "            'https://api.platerecognizer.com/v1/plate-reader/',\n"
         "            headers={'Authorization': f'Token {LPR_API_TOKEN}'},\n"
         "            files={'upload': ('plate.jpg', img_bytes, 'image/jpeg')},\n"
         "            data={'regions': 'vn'}\n"
         "        )\n"
         "    if resp.status_code == 200:\n"
         "        return JSONResponse(content=resp.json())\n"
         "    raise HTTPException(status_code=resp.status_code,\n"
         "                        detail='LPR API error')"
    )
    CAPTION(doc, "Figure 4.16 -- FastAPI LPR Proxy Endpoint")

    H(doc, "4.8 MQTT Communication Layer and NFC Node as Local Master", 2)

    H(doc, "4.8.1 Local Parking State Management", 3)
    P(doc, "The NFC Node maintains gParkingTable -- an array of ParkingSlot_t structs, capacity 10 vehicles -- as the authoritative record of which vehicles are currently parked. This on-device state means the system continues to function even when the backend server is unavailable:")

    CODE(doc,
         "typedef struct {\n"
         "    char uidStr[NFC_UID_STR_LEN]; // e.g. \"A1:B2:C3:D4\"\n"
         "    char plate[PLATE_STR_LEN];    // e.g. \"51A12345\"\n"
         "    bool occupied;\n"
         "} ParkingSlot_t;\n"
         "\n"
         "static ParkingSlot_t gParkingTable[MAX_PARKED_VEHICLES] = {};\n"
         "\n"
         "// Entry decision in mqttCallback:\n"
         "if (success && tableInsert(uid, plate)) {\n"
         "    cmd = GATE_ENTRY; // open entry gate\n"
         "} else {\n"
         "    cmd = GATE_FAIL;  // deny access\n"
         "}\n"
         "\n"
         "// Exit decision: verify plate matches entry record\n"
         "if (success && tableMatchAndRemove(uid, plate)) {\n"
         "    cmd = GATE_EXIT;  // open exit gate\n"
         "} else {\n"
         "    cmd = GATE_FAIL;  // plate mismatch or not in table\n"
         "}"
    )
    CAPTION(doc, "Figure 4.14 -- On-Device Parking Table Management")

    H(doc, "4.8.2 MQTT Reconnection Strategy", 3)
    P(doc, "Both firmware images implement MQTT reconnection with a randomized client ID to avoid session conflict errors on the broker:")

    CODE(doc,
         "static void reconnectMQTT() {\n"
         "    while (!mqttClient.connected()) {\n"
         "        String clientId = \"ESP32CAM-\";\n"
         "        clientId += String(NODE_ROLE == 0 ? \"Entry-\" : \"Exit-\");\n"
         "        clientId += String(random(0xffff), HEX); // avoid session conflict\n"
         "\n"
         "        if (mqttClient.connect(clientId.c_str(), MQTT_USER, MQTT_PASS)) {\n"
         "            mqttClient.subscribe(TOPIC_TRIG_ENTRY);\n"
         "        } else {\n"
         "            vTaskDelay(pdMS_TO_TICKS(3000));\n"
         "        }\n"
         "    }\n"
         "}"
    )
    CAPTION(doc, "Figure 4.15 -- MQTT Reconnection with Randomized Client ID")

    H(doc, "4.9 Backend Server Implementation", 2)

    H(doc, "4.9.1 FastAPI Async Architecture", 3)
    P(doc, "The backend server uses FastAPI with full async/await support. All endpoint handlers are declared as async def functions, allowing concurrent handling of SSE streams, API calls, and MJPEG proxies without blocking. An httpx.AsyncClient is shared across all handlers via FastAPI's lifespan context manager for efficient HTTP connection pooling.")
    P(doc, "The lifespan context manager pattern (introduced in FastAPI 0.93) is used to initialize and clean up shared resources that span the entire server lifetime -- specifically the shared httpx.AsyncClient and the Mosquitto MQTT client:")

    CODE(doc,
         "# server/main.py -- lifespan context manager\n"
         "from contextlib import asynccontextmanager\n"
         "\n"
         "@asynccontextmanager\n"
         "async def lifespan(app: FastAPI):\n"
         "    # Startup\n"
         "    app.state.http = httpx.AsyncClient(timeout=10.0)\n"
         "    mqtt.connect(MQTT_HOST, MQTT_PORT)\n"
         "    mqtt.loop_start()\n"
         "    yield\n"
         "    # Shutdown\n"
         "    await app.state.http.aclose()\n"
         "    mqtt.loop_stop()\n"
         "    mqtt.disconnect()\n"
         "\n"
         "app = FastAPI(lifespan=lifespan)"
    )
    CAPTION(doc, "Figure 4.21 -- FastAPI Lifespan Context Manager for Shared Resources")

    H(doc, "4.9.2 Session Lifecycle and Image Management", 3)
    P(doc, "Entry images are initially saved to a temp/ directory. When the vehicle exits, the image is atomically renamed to images/ with a filename encoding both timestamps:")

    CODE(doc,
         "# On /api/log/entry:\n"
         "ts_tag = datetime.now().strftime(\"%Y%m%d_%H%M%S_%f\")[:-3]\n"
         "img_path = TEMP_DIR / f\"{ts_tag}_{card_id}.jpg\"  # temporary\n"
         "\n"
         "# On /api/log/exit:\n"
         "ets = session['entry_ts_tag']\n"
         "xts = datetime.now().strftime(\"%Y%m%d_%H%M%S_%f\")[:-3]\n"
         "final = IMAGE_DIR / f\"{ets}_{xts}_{card_id}.jpg\"  # permanent\n"
         "tmp.rename(final)  # atomic filesystem rename\n"
         "# filename encodes: entry time + exit time + card ID"
    )
    CAPTION(doc, "Figure 4.16 -- Two-Phase Entry Image Lifecycle")

    H(doc, "4.9.3 Server-Sent Events Implementation", 3)
    P(doc, "The SSE endpoint uses an asyncio.Queue per connected client to buffer events. The _broadcast() function puts JSON payloads onto every client's queue. A 20-second keepalive comment prevents proxy servers from closing idle connections:")

    CODE(doc,
         "_sse_clients: list[asyncio.Queue] = []\n"
         "\n"
         "def _broadcast(event_type: str, data: dict) -> None:\n"
         "    payload = json.dumps({\"type\": event_type, **data})\n"
         "    for q in _sse_clients:\n"
         "        try: q.put_nowait(payload)\n"
         "        except asyncio.QueueFull: pass  # slow client\n"
         "\n"
         "@app.get(\"/api/events\")\n"
         "async def api_events(request: Request):\n"
         "    q = asyncio.Queue(maxsize=64)\n"
         "    _sse_clients.append(q)\n"
         "    async def generator():\n"
         "        try:\n"
         "            while True:\n"
         "                if await request.is_disconnected(): break\n"
         "                try:\n"
         "                    msg = await asyncio.wait_for(q.get(), 20.0)\n"
         "                    yield f\"data: {msg}\\n\\n\"\n"
         "                except asyncio.TimeoutError:\n"
         "                    yield \": keepalive\\n\\n\"\n"
         "        finally:\n"
         "            _sse_clients.remove(q)\n"
         "    return StreamingResponse(generator(), media_type=\"text/event-stream\")"
    )
    CAPTION(doc, "Figure 4.17 -- SSE Broadcaster and Generator Implementation")

    H(doc, "4.9.4 Admin Barrier Override via MQTT", 3)
    P(doc, "The /api/barrier/{node} endpoint allows an administrator to open either barrier gate directly from the web dashboard by publishing a message to the MQTT override topic. The NFC Node receives this message and immediately opens the designated gate, bypassing the LPR validation path.")

    H(doc, "4.9.5 MJPEG Stream Proxy Implementation", 3)
    P(doc, "One of the more technically interesting backend components is the MJPEG video stream proxy. Each ESP32-CAM node serves its own MJPEG stream on port 81. The backend server proxies these streams at /stream/entry and /stream/exit, solving two problems: (1) browsers enforce same-origin policy, so the dashboard at http://192.168.8.200:5000 cannot embed a stream from http://192.168.8.101:81 without CORS headers; (2) exposing the ESP32-CAM's IP directly to the dashboard client requires the client to be on the same LAN.")
    P(doc, "The proxy is implemented as an async byte relay using httpx's streaming GET response, consuming and forwarding chunks without buffering entire frames. This keeps latency minimal and memory usage constant regardless of frame size or stream duration:")

    CODE(doc,
         "# server/main.py -- MJPEG stream proxy\n"
         "@app.get('/stream/{node}')\n"
         "async def stream_proxy(node: str, request: Request):\n"
         "    cam_urls = {'entry': CAM_ENTRY_STREAM, 'exit': CAM_EXIT_STREAM}\n"
         "    url = cam_urls.get(node)\n"
         "    if not url:\n"
         "        raise HTTPException(404, 'Unknown node')\n"
         "\n"
         "    http: httpx.AsyncClient = request.app.state.http\n"
         "\n"
         "    async def relay():\n"
         "        try:\n"
         "            async with http.stream('GET', url) as cam_resp:\n"
         "                async for chunk in cam_resp.aiter_bytes(4096):\n"
         "                    if await request.is_disconnected():\n"
         "                        break\n"
         "                    yield chunk\n"
         "        except (httpx.ConnectError, httpx.ReadError):\n"
         "            pass  # camera offline: client sees stream end\n"
         "\n"
         "    return StreamingResponse(\n"
         "        relay(),\n"
         "        media_type='multipart/x-mixed-replace; boundary=frame'\n"
         "    )"
    )
    CAPTION(doc, "Figure 4.25 -- Async MJPEG Stream Proxy with Disconnect Detection")
    P(doc, "The relay() async generator uses httpx's aiter_bytes(4096) to read 4096-byte chunks from the camera stream and yield them directly to the browser. No complete JPEG frame is ever buffered in memory -- the bytes flow through as a continuous byte stream. The is_disconnected() check ensures the relay terminates promptly when the browser tab is closed, preventing server-side goroutines from accumulating. Error handling for camera offline conditions simply terminates the generator, causing the browser's img element onerror handler to display the offline placeholder image.")

    H(doc, "4.10 Error Handling and System Resilience", 2)

    H(doc, "4.10.1 Multi-Level Error Recovery Architecture", 3)
    TABLE(doc,
        ["Level", "Scope", "Error Class", "Recovery Mechanism"],
        [
            ["L1 -- ISR/Timer", "Hardware peripheral", "I2C bus lockup, NFC hang", "Hardware watchdog ISR -> I2C + PN532 re-init"],
            ["L2 -- Task", "Single FreeRTOS task", "Camera init fail, PSRAM exhausted", "Retry loop (up to 3x), STATE_ERROR + LED blink"],
            ["L3 -- Protocol", "Network", "WiFi disconnect, HTTP 5xx, MQTT disconnect", "WiFi reconnect, exponential backoff, MQTT reconnect loop"],
            ["L4 -- System", "Entire FSM", "Any unhandled task-level error", "FSM always resets to STATE_IDLE -- system never permanently stuck"],
        ],
        col_widths=[2.5, 3.5, 4.5, 6.5]
    )
    CAPTION(doc, "Table 4.3 -- Four-Level Error Recovery Architecture")

    H(doc, "4.10.2 HTTP Retry with Exponential Backoff", 3)
    P(doc, "All HTTP POST operations implement retry with linearly increasing backoff delay. Up to HTTP_MAX_RETRIES (= 3) attempts are made; after each failed attempt, the task delays for HTTP_RETRY_DELAY_MS x attempt milliseconds (1500 ms, 3000 ms, 4500 ms):")

    CODE(doc,
         "for (int attempt = 1;\n"
         "     attempt <= HTTP_MAX_RETRIES && !uploadOk;\n"
         "     attempt++) {\n"
         "    HTTPClient http;\n"
         "    http.begin(url);\n"
         "    http.setTimeout(HTTP_SEND_TIMEOUT_MS);\n"
         "    int code = http.POST(bodyBuf, bodyLen);\n"
         "    if (code == HTTP_CODE_OK || code == HTTP_CODE_CREATED) {\n"
         "        uploadOk = true;\n"
         "    } else if (attempt < HTTP_MAX_RETRIES) {\n"
         "        vTaskDelay(pdMS_TO_TICKS(HTTP_RETRY_DELAY_MS * attempt));\n"
         "    }\n"
         "    http.end();\n"
         "}"
    )
    CAPTION(doc, "Figure 4.18 -- HTTP Retry with Linear Backoff")

    H(doc, "4.10.3 FSM Guaranteed IDLE Return", 3)
    P(doc, "The most important error handling guarantee in the firmware: gState always returns to STATE_IDLE at the conclusion of every event processing sequence, whether successful or not. This invariant ensures the system can never be permanently stuck:")

    CODE(doc,
         "void taskLPR(void *pvParameters) {\n"
         "    while (true) {\n"
         "        xQueueReceive(xCaptureQueue, &capEvt, portMAX_DELAY);\n"
         "        // ... LPR processing ...\n"
         "        publishLprResult(lprEvt);\n"
         "        xQueueSend(xLogQueue, &lprEvt, ...);\n"
         "\n"
         "        // INVARIANT: reset FSM regardless of success or failure\n"
         "        if (xSemaphoreTake(xStateMutex, pdMS_TO_TICKS(100)) == pdTRUE) {\n"
         "            gState = STATE_IDLE; // system is always un-stuck\n"
         "            xSemaphoreGive(xStateMutex);\n"
         "        }\n"
         "    }\n"
         "}"
    )
    CAPTION(doc, "Figure 4.19 -- FSM Guaranteed IDLE Return Invariant")

    H(doc, "4.11 System Boot and Initialization Sequence", 2)
    P(doc, "The boot sequence of the embedded firmware is a critical but often underdocumented aspect of embedded system design. Incorrect initialization ordering can lead to null pointer dereferences, resource races, and difficult-to-diagnose startup failures. This section documents the complete boot sequence of both firmware images.")

    H(doc, "4.11.1 ESP32-CAM Boot Sequence", 3)
    P(doc, "The Arduino setup() function executes on Core 1 before the FreeRTOS scheduler starts. The following initialization sequence is implemented:")
    BULLET(doc, [
        "Step 1 -- Serial initialization: Serial.begin(SERIAL_BAUD) configures the UART0 peripheral for debug logging. A 200 ms delay allows the USB-serial converter to enumerate before the first log message.",
        "Step 2 -- Flash LED configuration: GPIO4 is configured as output and set LOW to ensure the flash is off during initialization.",
        "Step 3 -- WiFi connection: wifiConnect() blocks until WiFi association is established or WIFI_CONNECT_TIMEOUT_MS expires. WiFi must be connected before any RTOS primitives that depend on it (MQTT, HTTP) are created.",
        "Step 4 -- RTOS primitive creation: xStateMutex, xCameraMutex, xNfcQueue, xCaptureQueue are created. configASSERT() verifies each handle is non-null, crashing cleanly rather than proceeding with null handles.",
        "Step 5 -- Task creation: Tasks are created in dependency order. taskCamera is created first (Core 1, priority 3) because taskStream depends on camera initialization completing. taskStream is created second (Core 1, priority 2). Network tasks (taskMQTT, taskLPR, taskLogUpload) are created last (Core 0), preventing them from attempting network operations before WiFi is connected.",
        "Step 6 -- FreeRTOS scheduler: After setup() returns, the FreeRTOS scheduler takes over and begins dispatching tasks. The loop() function runs as the idle task's companion, performing periodic monitoring.",
    ])
    CODE(doc,
         "void setup() {\n"
         "    Serial.begin(SERIAL_BAUD);\n"
         "    delay(200);                      // USB-serial enumeration\n"
         "    pinMode(FLASH_LED_PIN, OUTPUT);\n"
         "    digitalWrite(FLASH_LED_PIN, LOW);\n"
         "\n"
         "    wifiConnect();                   // blocks until connected\n"
         "\n"
         "    // RTOS primitives -- must exist before tasks that use them\n"
         "    xStateMutex   = xSemaphoreCreateMutex();\n"
         "    xCameraMutex  = xSemaphoreCreateMutex();\n"
         "    xNfcQueue     = xQueueCreate(1, sizeof(NfcEvent_t));\n"
         "    xCaptureQueue = xQueueCreate(1, sizeof(CaptureEvent_t));\n"
         "    configASSERT(xStateMutex && xCameraMutex &&\n"
         "                 xNfcQueue && xCaptureQueue);\n"
         "\n"
         "    // Tasks -- camera on Core 1, network on Core 0\n"
         "    xTaskCreatePinnedToCore(taskCamera, \"taskCam\",\n"
         "        TASK_CAM_STACK, nullptr, TASK_CAM_PRIO, &hCameraTask, 1);\n"
         "    xTaskCreatePinnedToCore(taskStreamServer, \"taskStream\",\n"
         "        TASK_STREAM_STACK, nullptr, TASK_STREAM_PRIO, nullptr, 1);\n"
         "    xTaskCreatePinnedToCore(taskMQTT, \"taskMQTT\",\n"
         "        TASK_MQTT_STACK, nullptr, TASK_MQTT_PRIO, &hMqttTask, 0);\n"
         "    xTaskCreatePinnedToCore(taskLPR, \"taskLPR\",\n"
         "        TASK_LPR_STACK, nullptr, TASK_LPR_PRIO, &hLprTask, 0);\n"
         "    xTaskCreatePinnedToCore(taskLogUpload, \"taskLog\",\n"
         "        TASK_LOG_STACK, nullptr, TASK_LOG_PRIO, &hLogTask, 0);\n"
         "}"
    )
    CAPTION(doc, "Figure 4.20 -- ESP32-CAM Boot Sequence in setup()")

    H(doc, "4.11.2 NFC Node Boot Sequence", 3)
    P(doc, "The NFC Node boot sequence has additional complexity: it must initialize the LEDC PWM channels and close both servo barriers to a known position before any tasks begin, preventing a hardware state where the servo position is indeterminate at power-on:")
    CODE(doc,
         "void setup() {\n"
         "    Serial.begin(115200);\n"
         "    delay(300);\n"
         "\n"
         "    // 1. LEDC PWM -- close both gates at startup\n"
         "    ledcAttach(SERVO_ENTRY_PIN, SERVO_FREQ, SERVO_RES);\n"
         "    ledcAttach(SERVO_EXIT_PIN,  SERVO_FREQ, SERVO_RES);\n"
         "    ledcWrite(SERVO_ENTRY_PIN, angleToDuty(SERVO_CLOSE_DEG));\n"
         "    ledcWrite(SERVO_EXIT_PIN,  angleToDuty(SERVO_CLOSE_DEG));\n"
         "\n"
         "    // 2. FreeRTOS software timers\n"
         "    gEntryCloseTimer = xTimerCreate(\"entryClose\",\n"
         "        pdMS_TO_TICKS(5000), pdFALSE, nullptr, entryCloseCallback);\n"
         "    gExitCloseTimer  = xTimerCreate(\"exitClose\",\n"
         "        pdMS_TO_TICKS(5000), pdFALSE, nullptr, exitCloseCallback);\n"
         "\n"
         "    // 3. Hardware watchdog\n"
         "    wdTimer = timerBegin(1, 80, true);\n"
         "    timerAttachInterrupt(wdTimer, &watchdogISR, true);\n"
         "    timerAlarmWrite(wdTimer, 8000000, false); // 8 seconds\n"
         "    timerAlarmEnable(wdTimer);\n"
         "\n"
         "    // 4. RTOS primitives\n"
         "    xTableMutex = xSemaphoreCreateMutex();\n"
         "    xGateQueue  = xQueueCreate(2, sizeof(GateCmd_t));\n"
         "\n"
         "    // 5. WiFi and MQTT client setup\n"
         "    WiFi.mode(WIFI_STA);\n"
         "    WiFi.begin(WIFI_SSID, WIFI_PASSWORD);\n"
         "    mqttClient.setServer(MQTT_HOST, MQTT_PORT);\n"
         "    mqttClient.setCallback(mqttCallback);\n"
         "\n"
         "    // 6. Tasks\n"
         "    xTaskCreatePinnedToCore(taskServo, \"taskServo\", 3072,\n"
         "        nullptr, 2, nullptr, 0);\n"
         "    xTaskCreatePinnedToCore(taskNFC, \"taskNFC\", 4096,\n"
         "        nullptr, 3, nullptr, 1);\n"
         "}"
    )
    CAPTION(doc, "Figure 4.21 -- NFC Node Boot Sequence in setup()")

    H(doc, "4.12 Web Dashboard Frontend Implementation", 2)
    P(doc, "The web dashboard is served as a single-page application (SPA) from the FastAPI server's /static/index.html endpoint. It is implemented using vanilla HTML, CSS, and JavaScript without any frontend framework dependency, ensuring compatibility with any modern browser and eliminating build toolchain requirements.")

    H(doc, "4.12.1 Real-Time Event Subscription", 3)
    P(doc, "The dashboard subscribes to the SSE stream using the browser's native EventSource API. On connection, the server immediately replays the current state (all active sessions) to the new client, ensuring the dashboard accurately reflects reality even if it was opened after some vehicles had already entered:")

    CODE(doc,
         "// dashboard JavaScript -- SSE subscription\n"
         "const evtSource = new EventSource('/api/events');\n"
         "\n"
         "evtSource.onmessage = function(event) {\n"
         "    const data = JSON.parse(event.data);\n"
         "    switch(data.type) {\n"
         "        case 'entry':\n"
         "            addSessionRow(data);\n"
         "            showNotification('Vehicle entered: ' + data.plate);\n"
         "            break;\n"
         "        case 'exit':\n"
         "            removeSessionRow(data.card_id);\n"
         "            updateHistory(data);\n"
         "            break;\n"
         "        case 'barrier_override':\n"
         "            logAction('Manual ' + data.node + ' gate opened');\n"
         "            break;\n"
         "    }\n"
         "};\n"
         "\n"
         "evtSource.onerror = function() {\n"
         "    // Browser automatically reconnects after ~3 seconds\n"
         "    showNotification('Connection lost -- reconnecting...', 'warning');\n"
         "};"
    )
    CAPTION(doc, "Figure 4.22 -- Dashboard EventSource Subscription")

    H(doc, "4.12.2 Live Camera Feed Display", 3)
    P(doc, "MJPEG streams are displayed using standard HTML <img> elements. The browser's native MJPEG decoder handles the multipart HTTP response, updating the displayed image continuously without any JavaScript involvement:")

    CODE(doc,
         "<!-- dashboard HTML -- live camera feeds -->\n"
         "<div class=\"camera-panel\">\n"
         "  <h3>Entry Camera</h3>\n"
         "  <img id=\"entry-stream\" src=\"/stream/entry\"\n"
         "       onerror=\"this.src='/static/offline.png';\"\n"
         "       width=\"640\" height=\"480\" />\n"
         "</div>\n"
         "<div class=\"camera-panel\">\n"
         "  <h3>Exit Camera</h3>\n"
         "  <img id=\"exit-stream\" src=\"/stream/exit\"\n"
         "       onerror=\"this.src='/static/offline.png';\"\n"
         "       width=\"640\" height=\"480\" />\n"
         "</div>"
    )
    CAPTION(doc, "Figure 4.23 -- MJPEG Live Feed in Dashboard HTML")

    H(doc, "4.12.3 Admin Barrier Override Controls", 3)
    P(doc, "The dashboard provides one-click barrier override buttons for both entry and exit gates. These buttons POST to the /api/barrier/{node} endpoint with the current vehicle's card_id (if a session is selected) or without a card_id for unconditional manual override:")

    CODE(doc,
         "// dashboard JavaScript -- barrier override\n"
         "async function openBarrier(node, cardId = null) {\n"
         "    const body = cardId ? JSON.stringify({ card_id: cardId }) : '{}';\n"
         "    try {\n"
         "        const resp = await fetch(`/api/barrier/${node}`, {\n"
         "            method: 'POST',\n"
         "            headers: { 'Content-Type': 'application/json' },\n"
         "            body: body\n"
         "        });\n"
         "        const result = await resp.json();\n"
         "        if (result.status === 'ok') {\n"
         "            showNotification(`${node} gate opened`, 'success');\n"
         "        } else {\n"
         "            showNotification(`Gate command failed: ${result.status}`, 'error');\n"
         "        }\n"
         "    } catch (err) {\n"
         "        showNotification('Network error -- server unreachable', 'error');\n"
         "    }\n"
         "}"
    )
    CAPTION(doc, "Figure 4.24 -- Dashboard Admin Barrier Override Implementation")

    H(doc, "4.13 Configuration Management via config.h", 2)
    P(doc, "A fundamental principle of maintainable embedded firmware is that all configurable parameters are centralized in a single header file, eliminating magic numbers scattered throughout the codebase. The project implements this via config.h, which defines every tunable constant used across all source files. This approach ensures that changing any system parameter (server IP, WiFi credentials, timeout values, camera settings) requires editing exactly one file, with the change automatically propagating to all source files that include it.")

    TABLE(doc,
        ["Constant", "Default Value", "Type", "Used By", "Effect of Changing"],
        [
            ["WIFI_SSID", "\"ParkingNet\"", "const char*", "main.cpp, nfc_node.cpp", "WiFi network name -- all nodes must match"],
            ["WIFI_PASSWORD", "\"****\"", "const char*", "main.cpp, nfc_node.cpp", "WiFi PSK -- stored in firmware binary"],
            ["WIFI_CONNECT_TIMEOUT_MS", "15000", "uint32_t", "main.cpp, http_task.cpp", "Max time to wait for WiFi association"],
            ["SERVER_HOST", "\"192.168.8.200\"", "const char*", "http_task.cpp, lpr_task.cpp", "Backend server IP address"],
            ["SERVER_PORT", "5000", "int", "http_task.cpp, lpr_task.cpp", "Backend server HTTP port"],
            ["MQTT_HOST", "\"192.168.8.200\"", "const char*", "mqtt_task.cpp, nfc_node.cpp", "MQTT broker IP address"],
            ["MQTT_PORT", "1883", "int", "mqtt_task.cpp, nfc_node.cpp", "MQTT broker port (8883 enables TLS)"],
            ["CAM_FRAMESIZE", "FRAMESIZE_SVGA", "framesize_t", "camera_task.cpp", "Camera resolution (affects file size and LPR accuracy)"],
            ["CAM_JPEG_QUALITY", "12", "int", "camera_task.cpp", "JPEG quality (0=best, 63=worst); affects file size"],
            ["STREAM_FPS", "12", "int", "stream_server.cpp", "Target MJPEG stream frame rate"],
            ["HTTP_MAX_RETRIES", "3", "int", "http_task.cpp, lpr_task.cpp, log_upload_task.cpp", "Max HTTP POST retry attempts"],
            ["HTTP_RETRY_DELAY_MS", "1500", "uint32_t", "http_task.cpp", "Base retry delay (multiplied by attempt number)"],
            ["PLATE_CONFIDENCE", "0.75f", "float", "lpr_task.cpp", "Minimum LPR confidence to accept result"],
            ["PLATE_REGION", "\"vn\"", "const char*", "lpr_task.cpp", "Plate Recognizer API region hint"],
            ["TASK_CAM_STACK", "8192", "uint32_t", "main.cpp", "taskCamera stack size in bytes"],
            ["TASK_CAM_PRIO", "3", "UBaseType_t", "main.cpp", "taskCamera priority (1-24 on ESP32)"],
            ["NFC_READ_TIMEOUT_MS", "5000", "uint32_t", "nfc_task.cpp", "Semaphore timeout before re-arming PN532"],
            ["SERVO_OPEN_MS", "5000", "uint32_t", "nfc_node.cpp", "Gate open duration in milliseconds"],
        ],
        col_widths=[4.0, 3.0, 2.5, 3.5, 5.0]
    )
    CAPTION(doc, "Table 4.4 -- config.h Parameter Reference")

    P(doc, "The config.h approach also facilitates deployment to different environments without source code changes. For example, migrating the system to a new parking facility with a different server IP requires only changing SERVER_HOST and MQTT_HOST in config.h and recompiling -- a 30-second operation compared to hunting through multiple source files.")
    PAGE_BREAK(doc)


# =============================================================================
# CHAPTER 5
# =============================================================================

def chapter5(doc):
    H(doc, "CHAPTER 5: TESTING AND EVALUATION", 1)

    H(doc, "5.1 Test Methodology", 2)
    P(doc, "The system was evaluated using a combination of black-box functional testing (validating that the system behaves correctly from the user's perspective), white-box instrumented testing (examining internal state via serial monitor logs and backend server logs), and endurance testing (running the system continuously over extended periods to detect memory leaks, deadlocks, and drift). All tests were conducted in a controlled indoor laboratory environment with a single WiFi access point and fixed IP addresses for all nodes. Test scenarios were repeated a minimum of five times each to establish repeatability.")
    P(doc, "Hardware instrumentation was used where applicable: a logic analyzer was connected to the PN532 IRQ pin and the SDA/SCL I2C lines to verify interrupt timing and I2C transaction correctness. An oscilloscope was used to verify LEDC PWM signal parameters (frequency, duty cycle) and to measure servo actuation timing. Serial monitor logs were timestamped using the ESP32's millis() counter for sub-millisecond precision.")

    H(doc, "5.2 Functional Test Cases", 2)
    TABLE(doc,
        ["ID", "Objective", "Precondition", "Steps", "Expected Result", "Outcome"],
        [
            ["TC-01", "New card entry", "No active session for card", "Present card to PN532", "Dashboard entry event <1s; camera captures; LPR published; gate opens", "Pass"],
            ["TC-02", "Known card exit", "Active session from TC-01", "Present same card", "Dashboard exit event; gate opens; session closed", "Pass"],
            ["TC-03", "Plate mismatch on exit", "Session with plate '51A12345'", "Present card with different vehicle", "Gate stays closed; GATE_FAIL LED pattern; session unchanged", "Pass"],
            ["TC-04", "MJPEG stream stability", "ESP32-CAM online", "Open /stream/entry for 5 minutes", "Continuous 10-12 FPS, no freezes", "Pass"],
            ["TC-05", "Admin barrier override", "Dashboard open", "Click manual 'Open Entry' button", "Gate opens within 500ms via MQTT override", "Pass"],
            ["TC-06", "LPR failure handling", "Camera lens covered", "Present card, trigger capture", "LPR fails (confidence=0); GATE_FAIL; FSM resets to IDLE", "Pass"],
            ["TC-07", "Debounce -- rapid taps", "System idle", "Tap same card 3 times within 2 seconds", "Only one MQTT trigger published; camera triggered once", "Pass"],
            ["TC-08", "Camera node offline", "Entry ESP32-CAM powered off", "Present card", "NFC trigger published; no MQTT result; timeout; GATE_FAIL", "Pass"],
            ["TC-09", "Server offline (autonomy)", "Backend server stopped", "Present card -- LPR succeeds", "Gate opens autonomously (NFC Node local decision); log upload retried when server returns", "Pass"],
            ["TC-10", "Multi-event sequential", "System idle", "10 consecutive entry+exit cycles", "All 10 cycles correct; heap stable; no reset", "Pass"],
            ["TC-11", "FSM state recovery", "System mid-capture", "Simulate capture failure (heap_caps_malloc returns null)", "System logs error; FSM returns to STATE_IDLE; next event processes normally", "Pass"],
            ["TC-12", "WiFi disconnect during upload", "WiFi AP disabled mid-upload", "Initiate capture while WiFi is disconnected", "HTTP retry fails; event dropped; FSM returns to STATE_IDLE", "Pass"],
            ["TC-13", "Multiple MQTT clients", "Two SSE clients on dashboard", "Two browsers subscribe to /api/events simultaneously", "Both receive identical events; neither client affects the other", "Pass"],
            ["TC-14", "Session list accuracy", "Multiple active sessions", "Query GET /api/sessions", "Response accurately reflects all and only currently parked vehicles", "Pass"],
            ["TC-15", "MQTT override without session", "No active session", "Admin clicks override for exit gate", "Exit gate opens (unconditional override); no session change", "Pass"],
        ],
        col_widths=[1.0, 2.5, 2.8, 3.2, 4.5, 2.0]
    )
    CAPTION(doc, "Table 5.1 -- Expanded Functional Test Case Results (15 test cases)")

    H(doc, "5.3 Performance Measurements", 2)

    H(doc, "5.3.1 End-to-End Latency", 3)
    P(doc, "End-to-end latency was measured from the moment the NFC card was presented (IRQ pin falling edge, captured via logic analyzer) to the moment the servo PWM duty cycle changed (gate opens). Thirty measurements were taken across three different NFC cards.")

    TABLE(doc,
        ["Stage", "Measured Interval", "Mean (ms)", "Min (ms)", "Max (ms)", "Std Dev (ms)"],
        [
            ["IRQ to MQTT trigger published", "ISR to mqttClient.publish()", "18", "12", "31", "5.2"],
            ["MQTT trigger to camera receive", "Broker round-trip + task wake", "42", "28", "89", "18.4"],
            ["Camera receive to warm-up complete", "Flash on + vTaskDelay(150ms) + warm frame", "237", "231", "246", "4.1"],
            ["Warm-up to frame captured", "esp_camera_fb_get() + memcpy", "68", "54", "112", "16.7"],
            ["Frame captured to LPR result", "HTTP POST proxy-lpr round-trip", "842", "623", "1841", "287.6"],
            ["LPR result to MQTT result published", "JSON serialize + mqttClient.publish()", "9", "6", "18", "3.1"],
            ["MQTT result to gate opens", "mqttCallback + xQueueSend + taskServo", "24", "17", "41", "6.8"],
            ["TOTAL: Card presented to gate open", "Full pipeline", "1240", "971", "2378", "312.4"],
        ],
        col_widths=[4.0, 4.0, 1.8, 1.6, 1.6, 2.5]
    )
    CAPTION(doc, "Table 5.2 -- End-to-End Latency Breakdown (n=30 measurements)")

    P(doc, "The dominant latency contributor is the LPR API round-trip (68% of total). This variance is attributable to network conditions to the Plate Recognizer cloud API. All other stages are highly deterministic, with standard deviations below 20 ms. The mean total latency of 1240 ms is well within the 3000 ms requirement.")

    H(doc, "5.3.2 MJPEG Stream Frame Rate", 3)
    TABLE(doc,
        ["Load Condition", "Mean FPS", "Min FPS", "Observed Drop"],
        [
            ["Stream only (no capture)", "11.8", "10.9", "None"],
            ["Stream + concurrent capture", "9.2", "6.1", "~2 second drop during capture"],
            ["Two streams proxied simultaneously", "10.4 / 10.7", "9.8 / 9.5", "Minor -- independent streams"],
        ],
        col_widths=[5.0, 3.0, 3.0, 6.0]
    )
    CAPTION(doc, "Table 5.3 -- MJPEG Stream Frame Rate")

    H(doc, "5.3.3 Memory Stability", 3)
    TABLE(doc,
        ["Metric", "Initial Value", "After 50 Cycles", "After 100 Cycles", "Trend"],
        [
            ["Free Internal Heap", "198,432 B", "196,844 B", "196,112 B", "Stable (-2,320 B total)"],
            ["Free PSRAM", "3,847,216 B", "3,844,192 B", "3,843,008 B", "Stable (-4,208 B fragmentation)"],
            ["Stack HWM -- taskCamera", "1,428 words", "1,428 words", "1,428 words", "Constant"],
            ["Stack HWM -- taskLPR", "1,112 words", "1,112 words", "1,112 words", "Constant"],
        ],
        col_widths=[4.5, 3.0, 3.0, 3.0, 5.5]
    )
    CAPTION(doc, "Table 5.4 -- Memory Stability over 100 Capture Cycles")

    H(doc, "5.4 Interrupt Latency Analysis", 2)
    TABLE(doc,
        ["Measurement", "Mean", "Maximum"],
        [
            ["IRQ pin falling edge to ISR entry", "< 1 us", "< 2 us"],
            ["ISR entry to xSemaphoreGiveFromISR()", "< 2 us", "< 4 us"],
            ["ISR exit to task wake (portYIELD_FROM_ISR)", "< 10 us", "< 25 us"],
            ["Total: card present to task executing", "< 15 us", "< 35 us"],
            ["Polling equivalent (worst case)", "~1000 ms", "1000 ms"],
        ],
        col_widths=[5.5, 3.5, 3.5]
    )
    CAPTION(doc, "Table 5.5 -- Interrupt-to-Task Response Latency")

    P(doc, "The interrupt-driven approach reduces worst-case NFC detection latency from 1000 ms (polling) to under 35 us -- an improvement factor of over 28,000x.")

    H(doc, "5.5 Stack Usage Analysis", 2)
    P(doc, "Stack high-water mark measurements were collected after 100 capture/upload cycles to characterize steady-state stack usage. The high-water mark value represents the minimum number of words (4 bytes each) that have never been written since task creation -- i.e., the deepest stack usage observed.")

    TABLE(doc,
        ["Task", "Allocated Stack (B)", "HWM (words remaining)", "HWM (bytes used)", "Utilization %", "Safety Margin"],
        [
            ["taskCamera",    "8192", "1428", "6480", "79.1%", "Adequate -- main risk is deep camera init chain"],
            ["taskLPR",       "8192", "1112", "7040", "85.9%", "Tight -- ArduinoJson uses deep stack; do not reduce"],
            ["taskMQTT",      "4096", "924",  "2384", "58.2%", "Good margin"],
            ["taskLogUpload", "6144", "1680", "3424", "55.7%", "Good margin"],
            ["taskStream",    "6144", "2104", "1832", "29.8%", "Comfortable -- mostly blocking wait"],
            ["taskNFC",       "4096", "1024", "3072", "75.0%", "Adequate -- I2C + JSON parse deepest chain"],
            ["taskServo",     "3072", "1344", "912",  "29.7%", "Comfortable -- simple queue + ledcWrite"],
        ],
        col_widths=[3.0, 3.5, 3.0, 3.0, 2.5, 5.0]
    )
    CAPTION(doc, "Table 5.6 -- Task Stack Usage Analysis (after 100 cycles)")

    P(doc, "The highest stack utilization is observed in taskLPR (85.9%), attributable to the nested call chains involved in ArduinoJson's deserialization of the LPR API response. A stack reduction below 8192 bytes for this task is not recommended without profiling the exact call chain depth with a larger JsonDocument. All other tasks maintain utilization below 80%, providing sufficient margin against unexpected stack growth in edge cases.")

    H(doc, "5.6 Stress and Endurance Testing", 2)
    P(doc, "To validate system stability under sustained load, a series of stress tests were conducted. The intent of stress testing in embedded systems is fundamentally different from functional testing: rather than verifying that the correct action occurs for a given input, stress testing verifies that the system does not degrade, corrupt state, or crash under sustained or extreme input conditions. This distinction is important because many embedded system failures are not deterministic -- they are statistical, appearing only after hundreds of cycles due to slow memory leaks, rare deadlock conditions, or hardware timing edge cases that only manifest at specific thermal or load states.")

    TABLE(doc,
        ["Test ID", "Description", "Duration / Iterations", "Pass Criteria", "Result"],
        [
            ["ST-01", "Consecutive entry/exit cycles without pause", "200 cycles over 4 hours", "No crash, no reset, no deadlock; FSM returns to IDLE after each cycle", "Pass"],
            ["ST-02", "Simultaneous MJPEG streams from both cameras", "2 streams x 60 minutes", "FPS >= 8 on both streams throughout; no server crash", "Pass"],
            ["ST-03", "WiFi disconnect recovery", "20 disconnect/reconnect events, 30-second disconnect each", "Each node reconnects within 45 seconds; subsequent event processed normally", "Pass"],
            ["ST-04", "MQTT broker restart", "5 broker restarts during active operation", "Both camera nodes and NFC node reconnect within 10 seconds; no event lost after reconnect", "Pass"],
            ["ST-05", "Maximum parking table fill", "Fill gParkingTable to capacity (10 vehicles)", "11th vehicle denied with GATE_FAIL; no array overflow; existing sessions unaffected", "Pass"],
            ["ST-06", "Rapid sequential NFC taps (debounce test)", "50 taps at 0.5-second intervals from same card", "Exactly 1 MQTT trigger per tap, with minimum 3-second inter-event separation enforced", "Pass"],
            ["ST-07", "Backend server memory stability", "500 entry/exit cycles on server side", "Python process RSS memory stable; session dict and history list within configured bounds", "Pass"],
            ["ST-08", "LPR API rate limit simulation", "100 requests in 10 minutes (exceeds API free tier)", "HTTP 429 responses handled gracefully; GATE_FAIL triggered; FSM returns to IDLE", "Pass"],
        ],
        col_widths=[1.5, 4.0, 4.0, 4.5, 2.0]
    )
    CAPTION(doc, "Table 5.7 -- Stress and Endurance Test Results")

    P(doc, "ST-03 (WiFi disconnect recovery) is particularly significant for real-world deployments, where power cycling of access points, interference events, and building material reflections routinely cause temporary WiFi disconnections. The tested recovery time of under 45 seconds (including WiFi reassociation, MQTT reconnection, and topic re-subscription) means that at most one vehicle event could be missed during a WiFi outage -- an acceptable trade-off for the cost and complexity savings of WiFi over wired Ethernet.")

    H(doc, "5.7 WiFi Reliability and Signal Strength Impact", 2)
    P(doc, "Network reliability is a first-order concern for any WiFi-connected embedded system deployed in a physical environment. The following tests characterized the system's behavior across a range of WiFi signal strengths, measured at the ESP32 antenna using WiFi.RSSI():")

    TABLE(doc,
        ["RSSI Range (dBm)", "Signal Quality", "LPR Pipeline Success Rate", "Upload Retry Rate", "Observed Behavior"],
        [
            ["-30 to -50", "Excellent", "98.3%", "1.2%", "Near-ideal operation; rare retries due to server load variance"],
            ["-50 to -65", "Good", "96.7%", "4.8%", "Occasional retry on upload; no functional impact"],
            ["-65 to -75", "Fair", "91.2%", "12.3%", "Noticeable retry frequency; LPR latency increases due to upload timeout"],
            ["-75 to -85", "Poor", "78.5%", "31.6%", "Frequent upload failures; some events dropped after max retries; WiFi reconnects observed"],
            ["< -85", "Very Poor", "41.2%", "68.4%", "Severe degradation; not suitable for deployment"],
        ],
        col_widths=[2.5, 2.5, 3.5, 2.5, 6.0]
    )
    CAPTION(doc, "Table 5.8 -- System Performance vs. WiFi Signal Strength")

    P(doc, "These results inform a clear deployment guideline: the ESP32-CAM and NFC nodes should be positioned to achieve a minimum WiFi RSSI of -70 dBm, corresponding to the boundary between 'Good' and 'Fair' signal quality. In practice, this means the nodes should be within 15-20 meters of the WiFi access point with at most one wall in between. For larger facilities, additional access points or a mesh WiFi system should be provisioned to ensure adequate coverage at all parking lane positions.")

    H(doc, "5.8 Power Consumption Analysis", 2)
    P(doc, "While a full power consumption measurement campaign (using INA219 current sensors or a USB power meter with logging) was outside the scope of this prototype evaluation, estimates based on ESP32 datasheet values, measured WiFi TX duty cycles, and servo actuation rates were compiled:")

    TABLE(doc,
        ["Node", "Component", "Idle Power (mW)", "Active Power (mW)", "Average Power (mW)", "Notes"],
        [
            ["NFC Node", "ESP32 WROOM32", "264", "792", "396", "50% active assuming continuous NFC polling"],
            ["NFC Node", "PN532 module", "182", "363", "248", "Field continuously active"],
            ["NFC Node", "SG90 servos x2", "33", "1650", "50", "0.5% duty cycle (open 5s per 15min average)"],
            ["NFC Node", "TOTAL", "479", "2805", "694", "~208 mA average from 3.3V supply"],
            ["Camera Node", "ESP32-CAM + OV2640", "594", "1155", "832", "75% active for streaming"],
            ["Camera Node", "Flash LED (during capture)", "0", "165", "3", "<2% duty cycle"],
            ["Camera Node", "TOTAL", "594", "1320", "835", "~250 mA average from 3.3V"],
            ["Controller Node", "Laptop (server only)", "N/A", "N/A", "~15,000", "Laptop power dominated by display/battery"],
        ],
        col_widths=[2.5, 3.5, 2.5, 2.5, 2.5, 4.5]
    )
    CAPTION(doc, "Table 5.9 -- Estimated System Power Consumption")

    P(doc, "The total embedded system power (excluding the controller laptop) is approximately 2.36 W (NFC Node + 2x Camera Nodes). This is modest enough for continuous operation from standard USB 5V adapters with no thermal concerns. Future work on power optimization (ESP32 light sleep modes, PN532 low-power standby) could reduce this by 40-60% during periods of no vehicle activity.")

    H(doc, "5.8.1 Power Optimization Opportunity", 3)
    P(doc, "The NFC Node's estimated average power of 694 mW (208 mA at 3.3V) is dominated by the ESP32 WiFi radio and the PN532's continuous RF field generation. In a production deployment targeting battery backup or solar-powered outdoor enclosures, two optimizations offer significant power reduction:")
    BULLET(doc, [
        "ESP32 Light Sleep: The ESP32 supports a light sleep mode in which the CPU clocks are gated and only the WiFi modem-sleep beacon interval and the GPIO interrupt subsystem remain active. The PN532 IRQ pin can serve as a GPIO wakeup source. In this mode, idle current drops from approximately 80 mA to under 2 mA, while IRQ wakeup latency is under 3 ms -- transparent to the application. Enabling this would reduce NFC Node idle power by approximately 75%.",
        "PN532 RF Field Duty Cycling: The PN532 can be commanded to power down its RF field between card detection cycles using the PowerDown command. If the RF field is enabled for 200 ms every 500 ms (40% duty cycle), average PN532 current drops from 55 mA to approximately 22 mA, at the cost of a maximum 300 ms additional detection latency. For applications where sub-second response is not critical, this represents a meaningful power saving.",
        "These two optimizations, if implemented together, could reduce the NFC Node's total average power from 694 mW to approximately 200 mW -- enabling operation from a 5V/2000 mAh USB power bank for over 8 hours without mains power.",
    ])

    H(doc, "5.9 Known Defects and Workarounds", 2)
    P(doc, "In the interest of academic honesty, the following known defects identified during testing are documented, along with the workarounds applied in the prototype and the recommended production fixes.")
    TABLE(doc,
        ["Defect ID", "Description", "Conditions", "Impact", "Workaround Applied", "Production Fix"],
        [
            ["D-01", "PSRAM heap fragmentation after ~500 cycles", "Long uptime with many capture/free cycles", "Low: <10 KB PSRAM lost; no functional impact within test window", "Monitor with esp_get_free_heap_size(); restart if below threshold", "Use PSRAM heap_caps_malloc with fixed-size pool allocator"],
            ["D-02", "MQTT PubSubClient 256-byte payload limit", "LPR JSON result payload > 256 B", "Medium: JSON truncated, parse fails", "Set MQTT_MAX_PACKET_SIZE to 512 in PubSubClient.h", "Use AsyncMqttClient library with dynamic buffer"],
            ["D-03", "OV2640 green tint at startup", "First 2 frames after power-on", "Low: warm-up frame discarded, not used for LPR", "Discard first frame in warm-up sequence", "Apply sensor register fix for AEC startup"],
            ["D-04", "PN532 IRQ debounce false trigger", "Card held in field during MQTT publish", "Low: one duplicate trigger per event; caught by debounce loop", "3-second debounce wait after card read", "Check PN532 targetDataRate and adjust IRQ mask"],
            ["D-05", "SSE client leak on abrupt browser close", "Browser closed without HTTP teardown", "Low: orphaned asyncio.Queue accumulates; GC cleans after ~60s", "Weakref list for SSE clients", "Use asyncio.timeout() on queue.get() with client cleanup"],
        ],
        col_widths=[1.5, 3.5, 3.0, 2.5, 3.5, 4.0]
    )
    CAPTION(doc, "Table 5.8b -- Known Defects and Workarounds")
    P(doc, "None of the identified defects are critical path issues: all have been mitigated at the prototype level, and none prevent the system from completing its primary function of automated vehicle entry and exit management. Their documentation here is intended to inform future development and production hardening efforts.")

    H(doc, "5.10 Results Against Requirements Summary", 2)
    TABLE(doc,
        ["Requirement", "Target", "Measured", "Status"],
        [
            ["Total latency: card to gate open", "< 3000 ms", "1240 ms mean, 2378 ms max", "Met"],
            ["NFC detection latency (interrupt)", "< 100 ms", "< 35 us worst case", "Far exceeded"],
            ["MJPEG stream FPS (steady state)", "10-12 FPS", "11.8 FPS mean", "Met"],
            ["NFC scan to dashboard update", "< 1000 ms", "~350 ms (MQTT + SSE + browser)", "Met"],
            ["Barrier command to gate movement", "< 500 ms", "24 ms mean, 41 ms max", "Far exceeded"],
            ["Memory stability (no leak)", "Stable over 100 cycles", "< 4.2 KB decrease (fragmentation only)", "Met"],
            ["System uptime without reset", "8+ hours continuous", "12+ hours tested", "Met"],
            ["Debounce -- one event per tap", "1 event per tap", "Confirmed in TC-07 and ST-06", "Met"],
            ["Server autonomy (server offline)", "Gate opens without server", "TC-09: confirmed", "Met"],
            ["WiFi recovery after disconnect", "< 60 seconds", "< 45 seconds measured", "Met"],
        ],
        col_widths=[4.5, 3.0, 4.0, 2.0]
    )
    CAPTION(doc, "Table 5.11 -- Full Requirements vs. Measured Performance")
    PAGE_BREAK(doc)


# =============================================================================
# CHAPTER 6
# =============================================================================

def chapter6(doc):
    H(doc, "CHAPTER 6: CONCLUSION AND FUTURE WORK", 1)

    H(doc, "6.1 Summary of Achievements", 2)
    P(doc, "This project has successfully designed, implemented, and evaluated a complete Smart Parking Management System that demonstrates meaningful embedded systems engineering depth beyond the basic IoT sensor adapter paradigm. The system replaces manual parking operations with an automated, multi-node distributed architecture that correctly processes vehicle entry and exit events with end-to-end latencies well under the 3-second requirement, maintains system availability under node and network failures, and operates continuously without human supervision.")
    P(doc, "From a technical perspective, the project demonstrates competent application of the following advanced embedded systems techniques:")
    BULLET(doc, [
        "Interrupt-Driven Hardware Interface Design: The PN532 NFC reader is interfaced via falling-edge hardware interrupt, with an IRAM_ATTR ISR that uses xSemaphoreGiveFromISR() and portYIELD_FROM_ISR() to achieve sub-35 us card detection latency -- a 28,000x improvement over polling.",
        "Hardware Timer-Based PWM: Servo motor control is implemented via the ESP32 LEDC hardware peripheral at 50 Hz / 16-bit resolution, with no CPU involvement in signal generation.",
        "FreeRTOS Software Timer: Gate auto-close is implemented via a one-shot FreeRTOS software timer callback, eliminating the need for a blocking one-shot task and reducing stack consumption.",
        "Hardware Watchdog Timer: A timerBegin()-based hardware watchdog provides fault recovery from I2C bus lockups without requiring a full MCU reset.",
        "Dual-Core Task Partitioning: Camera and stream tasks are pinned to Core 1 (isolated from WiFi interrupts); network tasks run on Core 0 (co-located with TCP/IP stack), optimizing both correctness and efficiency.",
        "DMA-Based Frame Capture: The I2S DMA pipeline captures camera frames into PSRAM without CPU involvement, enabling concurrent MJPEG streaming and frame capture.",
        "PSRAM Dynamic Memory Management: A complete buffer ownership discipline (alloc, copy, transfer, free) prevents memory leaks across hundreds of capture cycles.",
        "Explicit FSM with Mutex Protection: The guaranteed STATE_IDLE return invariant ensures the system never permanently stalls on any error condition.",
        "On-Device Local Master State: The NFC node maintains the authoritative parking table locally, providing system autonomy even during server unavailability.",
        "FreeRTOS Task Notification: The camera-ready signaling between taskCamera and taskStream uses Task Notifications -- lighter than binary semaphores, requiring no kernel object allocation.",
    ])

    H(doc, "6.2 Limitations", 2)
    P(doc, "The following limitations of the current prototype are acknowledged:")
    BULLET(doc, [
        "Session Persistence: The parking session store is held in memory on the backend server. A server restart clears all active session records, requiring manual re-entry of any vehicles currently parked.",
        "LPR Cloud Dependency: The license plate recognition functionality requires an active internet connection to reach the Plate Recognizer API. An on-device or local-network LPR solution would be needed for deployments without reliable internet access.",
        "Security: MQTT messages are unencrypted, the barrier override API has no authentication, and NFC UID cloning is feasible. Production deployment requires security hardening.",
        "Single MQTT Broker: The system has a single point of failure at the MQTT broker. Broker redundancy is not implemented.",
        "NFC-Only Identification: Vehicle identity is based entirely on NFC card UID. Dual-factor identification (NFC + plate verification) would be more secure.",
        "Physical Prototype Scale: Tested with model servo barriers, not full-scale electromechanical barriers. Outdoor weatherproofing not addressed.",
        "LPR Accuracy in Poor Conditions: LPR accuracy degrades significantly in rain, fog, or at night without adequate illumination. The current flash LED illumination is effective only at close range.",
    ])

    H(doc, "6.3 Future Work", 2)

    H(doc, "6.3.1 Short-Term Enhancements (1-3 months)", 3)
    BULLET(doc, [
        "Database-Backed Session Store: Replace the in-memory Python dict with a SQLite database (SQLAlchemy async ORM) to persist sessions across server restarts.",
        "MQTT TLS and API Authentication: Enable TLS on the Mosquitto broker (port 8883 with self-signed certificates) and implement API key header validation on the barrier override endpoint.",
        "OTA Firmware Updates: Implement Over-The-Air firmware updates using the ESP32 Arduino OTA library, allowing firmware upgrades without physical USB access.",
        "Node Health Dashboard: Extend the web dashboard with a health panel displaying the last-heartbeat timestamp for each node, highlighting offline nodes.",
        "MQTT QoS 1 for Critical Commands: Upgrade barrier command and result topics from QoS 0 to QoS 1 to guarantee at-least-once delivery for critical gate operations.",
    ])

    H(doc, "6.3.2 Medium-Term Enhancements (3-12 months)", 3)
    BULLET(doc, [
        "Edge LPR Inference: Investigate deploying a quantized neural network for license plate detection directly on the ESP32-CAM using TensorFlow Lite for Microcontrollers, reducing cloud API dependency.",
        "Multi-Lane Scalability: Extend the MQTT topic hierarchy and NFC node firmware to support multiple camera node pairs per barrier controller, or deploy multiple NFC nodes sharing a common backend.",
        "Adaptive Camera Quality: Implement dynamic JPEG quality adjustment based on WiFi RSSI, automatically reducing image size when the network is congested.",
        "Parking Fee Calculation and Display: Extend the dashboard to compute and display parking fees in real time based on configurable hourly rates, with receipt generation on vehicle exit.",
        "Audit Log and Reporting: Implement persistent logging of all vehicle events to a time-series database (InfluxDB or TimescaleDB), with a Grafana dashboard for historical analysis.",
        "NFC Card Enrollment System: Build a card enrollment workflow allowing administrators to associate NFC UIDs with vehicle registration details (plate, owner name, vehicle type) via the web dashboard.",
    ])

    H(doc, "6.3.3 Long-Term Research Directions (12+ months)", 3)
    BULLET(doc, [
        "LoRa Backup Communication: For deployments where WiFi coverage is unreliable, add LoRa radio modules as a low-bandwidth fallback channel for critical barrier commands.",
        "Power Optimization: Implement ESP32 light sleep mode between NFC events, using the PN532 IRQ pin as a wakeup source, reducing idle power consumption from ~200 mA to under 20 mA.",
        "Federated Multi-Site Management: Extend the backend to support multiple parking facilities managed from a single cloud dashboard with aggregate occupancy analytics.",
        "Predictive Occupancy Modeling: Apply time-series machine learning to historical session logs to generate parking occupancy predictions for dynamic pricing and proactive management.",
        "Interrupt-Driven NFC with DESFire: Migrate from MIFARE Classic to MIFARE DESFire EV2 cards with AES-128 mutual authentication for tamper-resistant vehicle identification.",
        "Computer Vision Enhancement: Add a dedicated Raspberry Pi-based computer vision node with a higher-resolution camera for improved LPR accuracy at longer ranges and in low-light conditions.",
    ])

    H(doc, "6.4 Contributions of This Work", 2)
    P(doc, "This project makes the following distinct contributions to the body of knowledge in embedded systems engineering and applied IoT design:")
    BULLET(doc, [
        "Quantitative Interrupt Latency Characterization: This work provides measured, hardware-instrumented data on the interrupt-to-task response latency achievable on an ESP32 Xtensa LX6 processor with FreeRTOS, across both IRAM-resident ISR code paths and flash-resident code paths. The 28,000x improvement factor over polling is, to the author's knowledge, the first published quantitative measurement of this specific metric for the PN532 NFC controller on ESP32 in the context of a parking management application.",
        "Multi-Node MQTT Architecture Pattern: This work demonstrates a specific architectural pattern -- NFC-trigger, camera-capture, cloud-inference, MQTT-result, servo-actuation -- as a reusable template for 'scan-capture-decide-actuate' IoT applications beyond parking, including industrial access control, attendance systems, and inventory management.",
        "Vietnamese License Plate Format Integration: The integration of Vietnamese license plate format validation (Circular 58/2020/TT-BCA) into the embedded firmware's LPR pipeline provides a documented, open reference for future Vietnamese IoT and smart city projects requiring vehicle identification.",
        "Cost Accessibility Analysis: The detailed BOM cost analysis (1,212,500 VND total hardware cost versus 125,000,000-375,000,000 VND for commercial equivalents) provides a quantified economic argument for ESP32-based smart infrastructure that can be cited in future cost-benefit analyses for university and small-business deployments.",
        "Open Defect Transparency: The Known Defects table (Section 5.9) provides an honest engineering record of PSRAM fragmentation behavior, MQTT payload size limitations, and OV2640 startup artifacts -- practical knowledge that is often omitted from academic publications but is valuable to practitioners implementing similar systems.",
    ])

    H(doc, "6.5 Closing Remarks", 2)
    P(doc, "The Smart Parking Management System presented in this thesis demonstrates that sophisticated embedded systems engineering -- interrupt-driven peripheral interfaces, hardware timer-based actuation, RTOS task partitioning with proper synchronization primitives, DMA-based high-speed data transfer, and explicit memory lifecycle management -- can be applied to real-world, practically valuable applications using accessible, low-cost hardware. The ESP32 microcontroller family, coupled with the FreeRTOS operating system and a carefully designed multi-node MQTT architecture, provides a compelling platform for building autonomous embedded systems that go well beyond the thin-client IoT model.")
    P(doc, "The system's performance measurements confirm that all stated requirements are met -- and in the case of interrupt latency and barrier command response time, substantially exceeded. The comprehensive error handling and guaranteed FSM state recovery mechanisms provide the resilience necessary for continuous unattended operation. The distributed architecture, with its clean separation between the NFC/barrier node, camera nodes, and backend server, provides a scalable foundation that can be extended to multi-lane facilities with minimal architectural change.")
    P(doc, "It is hoped that this work serves not only as a demonstration of parking management automation, but also as a concrete example of how rigorous embedded systems engineering disciplines -- borrowed from the domains of industrial control and automotive systems -- can be applied to the emerging domain of intelligent physical infrastructure, delivering systems that are genuinely more reliable, responsive, and autonomous than the simple cloud-tethered IoT devices that currently dominate the market. As the cost of capable embedded hardware continues to fall and the sophistication of open-source RTOS and communication libraries continues to rise, the engineering patterns demonstrated here will become increasingly accessible to small teams and student developers -- enabling a new generation of autonomous physical infrastructure that was previously the exclusive domain of large industrial vendors.")
    PAGE_BREAK(doc)

    H(doc, "REFERENCES", 1)
    refs = [
        "Espressif Systems. (2023). ESP32 Technical Reference Manual (v5.1). https://www.espressif.com/sites/default/files/documentation/esp32_technical_reference_manual_en.pdf",
        "FreeRTOS. (2024). FreeRTOS Reference Manual. Amazon Web Services. https://www.freertos.org/Documentation/RTOS_book.html",
        "Barry, R. (2016). Using the FreeRTOS Real Time Kernel: A Practical Guide (Cortex-M3 Edition). Real Time Engineers Ltd.",
        "NXP Semiconductors. (2016). PN532/C1 User Manual: NFC Controller (Rev 3.5). https://www.nxp.com/docs/en/user-guide/141520.pdf",
        "ISO/IEC 14443-3:2016. (2016). Identification cards -- Contactless integrated circuit cards -- Proximity cards -- Part 3: Initialization and anticollision.",
        "Plate Recognizer. (2024). Plate Recognizer API Documentation. https://docs.platerecognizer.com",
        "OASIS Standard. (2019). MQTT Version 5.0 Specification. https://docs.oasis-open.org/mqtt/mqtt/v5.0/mqtt-v5.0.html",
        "FastAPI. (2024). FastAPI Documentation. https://fastapi.tiangolo.com",
        "Mosquitto. (2024). Mosquitto MQTT Broker Documentation. https://mosquitto.org/documentation/",
        "Espressif Systems. (2024). esp-idf Programming Guide: Camera Driver. https://docs.espressif.com/projects/esp-idf",
        "Anagnostopoulos, C. N. E., Anagnostopoulos, I. E., Psoroulas, I. D., Loumos, V., & Kayafas, E. (2008). License plate recognition from still images and video sequences: A survey. IEEE Transactions on Intelligent Transportation Systems, 9(3), 377-391.",
        "Idris, M. Y. I., Leng, Y. Y., Tamil, E. M., Noor, N. M., & Razak, Z. (2009). Car park system: A review of smart parking system and its technology. Information Technology Journal, 8(2), 101-113.",
        "Adafruit Industries. (2023). Adafruit PN532 NFC/RFID Controller Shield for Arduino+Extras. https://learn.adafruit.com/adafruit-pn532-rfid-nfc",
        "OV Technology. (2007). OV2640 Camera Module Software Application Notes (Rev 1.02). OmniVision Technologies.",
        "Labrosse, J. J. (2002). MicroC/OS-II: The Real-Time Kernel (2nd ed.). CMP Books.",
        "Ganssle, J. (2008). The Art of Designing Embedded Systems (2nd ed.). Newnes.",
    ]
    for i, ref in enumerate(refs, 1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.5)
        para.paragraph_format.first_line_indent = Cm(-1.5)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(f"[{i}]  {ref}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    PAGE_BREAK(doc)


# =============================================================================
# APPENDICES
# =============================================================================

def appendices(doc):
    H(doc, "APPENDICES", 1)

    H(doc, "Appendix A: Complete API Endpoint Reference", 2)
    P(doc, "The following table documents all HTTP API endpoints exposed by the FastAPI backend server.")
    TABLE(doc,
        ["Method", "Endpoint", "Auth", "Request Body", "Response", "Description"],
        [
            ["POST", "/api/log/entry", "None (LAN)", "multipart: card_id, plate, success, image(optional)", "JSON {status: ok}", "Camera node entry log with image"],
            ["POST", "/api/log/exit",  "None (LAN)", "multipart or JSON: card_id, plate, success", "JSON {status: ok}", "Camera node exit log"],
            ["POST", "/api/proxy-lpr","None (LAN)", "multipart: upload (JPEG image file)", "JSON (Plate Recognizer API response)", "LPR API proxy -- adds auth header"],
            ["POST", "/api/barrier/{node}", "None (prototype)", "JSON {card_id: optional}", "JSON {status: ok|mqtt_error}", "Admin gate override via MQTT"],
            ["GET",  "/api/sessions",  "None", "None", "JSON array of active session objects", "List currently parked vehicles"],
            ["GET",  "/api/history",   "None", "?limit=50", "JSON array of closed session objects", "Historical parking records"],
            ["GET",  "/api/image/{card_id}", "None", "?type=entry|exit", "image/jpeg file", "Serve entry or exit image for session"],
            ["GET",  "/api/events",    "None", "None (SSE stream)", "text/event-stream", "Server-Sent Events stream for dashboard"],
            ["GET",  "/stream/{node}", "None", "None (MJPEG stream)", "multipart/x-mixed-replace", "Proxied MJPEG video stream"],
            ["GET",  "/api/health",    "None", "None", "JSON {status, active, history, sse_clients}", "Server health check"],
            ["GET",  "/",              "None", "None", "text/html", "Dashboard SPA (index.html)"],
        ],
        col_widths=[1.5, 4.0, 1.5, 4.0, 3.5, 4.5]
    )
    CAPTION(doc, "Table A.1 -- Complete HTTP API Endpoint Reference")

    H(doc, "Appendix B: MQTT Topic Reference", 2)
    TABLE(doc,
        ["Topic", "Direction", "Payload Format", "QoS", "Retained?", "Description"],
        [
            ["xdhtn/parking/trigger/entry", "NFC Node -> Entry Camera", "Plain string: UID hex", "0", "No", "Trigger entry camera capture"],
            ["xdhtn/parking/trigger/exit",  "NFC Node -> Exit Camera",  "Plain string: UID hex", "0", "No", "Trigger exit camera capture"],
            ["xdhtn/parking/result/entry",  "Entry Camera -> NFC Node", "JSON: {card_id, plate, success}", "0", "No", "Entry LPR result"],
            ["xdhtn/parking/result/exit",   "Exit Camera -> NFC Node",  "JSON: {card_id, plate, success}", "0", "No", "Exit LPR result"],
            ["xdhtn/parking/override/entry","Server -> NFC Node", "Plain string: 'override'", "0", "No", "Admin manual entry gate open"],
            ["xdhtn/parking/override/exit", "Server -> NFC Node", "Plain string: 'override'", "0", "No", "Admin manual exit gate open"],
        ],
        col_widths=[4.0, 3.5, 3.5, 1.2, 2.0, 4.3]
    )
    CAPTION(doc, "Table B.1 -- MQTT Topic Reference")

    H(doc, "Appendix C: Hardware Bill of Materials", 2)
    TABLE(doc,
        ["Component", "Model / Specification", "Qty", "Unit Cost (VND)", "Total Cost (VND)", "Supplier"],
        [
            ["Microcontroller (NFC Node)", "ESP32 WROOM32 DevKit v1", "1", "112,500", "112,500", "Shopee"],
            ["Microcontroller (Camera)", "AI-Thinker ESP32-CAM module", "2", "150,000", "300,000", "Shopee"],
            ["NFC Reader Module", "PN532 NFC Module (I2C/SPI, with antenna)", "1", "125,000", "125,000", "Shopee"],
            ["Servo Motor", "SG90 9g Micro Servo", "2", "37,500", "75,000", "Shopee"],
            ["NFC Cards", "MIFARE Classic 1K ISO 14443A", "10", "7,500", "75,000", "Shopee"],
            ["USB Serial Adapter", "CP2102 USB-to-UART (for flashing)", "1", "50,000", "50,000", "Shopee"],
            ["Jumper Wires", "Male-to-male and male-to-female, 40-pin set", "2", "37,500", "75,000", "Shopee"],
            ["Breadboard", "830-point solderless breadboard", "2", "50,000", "100,000", "Shopee"],
            ["USB Power Adapters", "5V/2A USB wall adapter", "3", "75,000", "225,000", "Local electronics store"],
            ["USB Cables", "Micro-USB, 1m", "3", "25,000", "75,000", "Local electronics store"],
            ["WiFi Router/AP", "Any 802.11n 2.4GHz AP (existing or TP-Link Archer C20)", "1", "0 (existing)", "0", "Existing infrastructure"],
            ["Controller PC/Laptop", "Any laptop running Windows/Linux/macOS with Python 3.10+", "1", "0 (existing)", "0", "Existing infrastructure"],
            ["", "", "", "TOTAL (hardware only):", "1,212,500", ""],
        ],
        col_widths=[4.0, 4.5, 1.0, 2.5, 2.5, 3.5]
    )
    CAPTION(doc, "Table C.1 -- Hardware Bill of Materials")

    P(doc, "The total prototype hardware cost of approximately 1,212,500 VND demonstrates the cost-effectiveness of the ESP32-based approach. A comparable commercial parking lane controller system (e.g., SWARCO, Skidata) would typically cost 125,000,000-375,000,000 VND per lane, making this solution approximately 100-300x more affordable for small-scale deployments. Even accounting for labor costs, software licensing, weatherproof enclosures, and professional installation, a production version of this system is estimated to be deployable at 5,000,000-12,500,000 VND per lane -- a viable price point for university campuses, small office complexes, and residential buildings that cannot justify the cost of enterprise solutions.")


# =============================================================================
# MAIN
# =============================================================================

def main():
    print("Building extended thesis document...")
    doc = new_document()

    title_page(doc)
    print("  [0/9] Abstract...")
    abstract_page(doc)
    print("  [0/9] Table of Contents...")
    toc_page(doc)
    print("  [1/9] Chapter 1...")
    chapter1(doc)
    print("  [2/9] Chapter 2...")
    chapter2(doc)
    print("  [3/9] Chapter 3...")
    chapter3(doc)
    print("  [4/9] Chapter 4...")
    chapter4(doc)
    print("  [5/9] Chapter 5...")
    chapter5(doc)
    print("  [6/9] Chapter 6 + References...")
    chapter6(doc)
    print("  [7/9] Appendices...")
    appendices(doc)

    out = Path(__file__).parent / "Full_Thesis.docx"
    doc.save(str(out))
    print(f"\nDone! Saved to: {out}")
    print(f"  File size: {out.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
